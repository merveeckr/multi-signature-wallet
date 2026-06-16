"""
MultiSigWallet — Merve İlerleme Raporu (DOCX üreteci)
Çalıştır: python scripts/generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Sayfa kenar boşlukları ────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Renkler ───────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x18, 0x5F, 0xA5)   # Merve
GREEN      = RGBColor(0x0F, 0x6E, 0x56)   # Emine
ORANGE     = RGBColor(0x85, 0x4F, 0x0B)   # İrem
PURPLE     = RGBColor(0x53, 0x4A, 0xB7)   # Ortak
DARK       = RGBColor(0x1A, 0x20, 0x2C)
MUTED      = RGBColor(0x71, 0x80, 0x96)
RED        = RGBColor(0x99, 0x1B, 0x1B)
DARK_GREEN = RGBColor(0x16, 0x65, 0x34)

# ── Yardımcı fonksiyonlar ─────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(text, level=1, color=DARK):
    p   = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color
    run.font.bold      = True
    run.font.size      = Pt({1: 16, 2: 13, 3: 11}.get(level, 11))
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_para(text="", bold=False, italic=False, size=10, color=DARK,
             space_after=4, indent=0):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold        = bold
    run.font.italic      = italic
    run.font.size        = Pt(size)
    run.font.color.rgb   = color
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Cm(indent)
    return p

def add_code(text, indent=0):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name        = "Courier New"
    run.font.size        = Pt(9)
    run.font.color.rgb   = RGBColor(0x1E, 0x29, 0x3B)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(indent if indent else 0.5)
    # Açık gri arka plan
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F1F5F9")
    pPr.append(shd)
    return p

def add_bullet(text, bold_prefix=None, color=DARK, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold      = True
        r.font.size      = Pt(10)
        r.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size      = Pt(10)
    r2.font.color.rgb = color
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_table(headers, rows, header_bg="185FA5", col_widths=None):
    t   = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Başlık satırı
    hdr_row = t.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        run  = cell.paragraphs[0].add_run(hdr)
        run.font.bold      = True
        run.font.size      = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment      = WD_ALIGN_VERTICAL.CENTER

    # Veri satırları
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, "F8FAFC")
            run  = cell.paragraphs[0].add_run(str(cell_text))
            run.font.size = Pt(9)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Sütun genişlikleri
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def divider(color_hex="E2E8F0"):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)

def label_box(text, color_rgb, bg_hex):
    """Renkli etiket kutusu — kızlara 'bu bölüm sana' demek için"""
    p   = doc.add_paragraph()
    run = p.add_run(f"  {text}  ")
    run.font.bold      = True
    run.font.size      = Pt(9)
    run.font.color.rgb = color_rgb
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  bg_hex)
    pPr.append(shd)
    p.paragraph_format.space_after = Pt(4)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# KAPAK
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Multi-Signature Wallet")
r.font.size      = Pt(22)
r.font.bold      = True
r.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Merve'nin Yaptıkları — Teknik İlerleme Raporu")
r2.font.size      = Pt(13)
r2.font.color.rgb = MUTED

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("COM 4532 · Blockchain · 25 Mayıs 2026")
r3.font.size      = Pt(10)
r3.font.color.rgb = MUTED

doc.add_paragraph()

# Renk legend
add_para("Bu raporda renk kodlaması:", bold=True, size=9, color=MUTED)
legend_items = [
    ("█ Merve Çakır — Smart Contract & Güvenlik", BLUE),
    ("█ Emine Binay — Web3.py Backend (sana özel notlar bu renkle)",  GREEN),
    ("█ İrem Keskin — Test & Frontend (sana özel notlar bu renkle)",   ORANGE),
]
for txt, clr in legend_items:
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.size      = Pt(9)
    r.font.color.rgb = clr
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM 0 — GENEL DURUM
# ══════════════════════════════════════════════════════════════════════════════

add_heading("0. Genel Durum", level=1, color=BLUE)
add_para(
    "Bu rapor Merve'nin proje boyunca Solidity tarafında yaptığı her şeyi "
    "belgeler. Emine ve İrem'in entegrasyon ve test süreçleri için bilmesi "
    "gereken her şey açıkça işaretlenmiştir.",
    size=10
)
doc.add_paragraph()

add_table(
    headers=["Görev", "Hafta", "Durum"],
    rows=[
        ["3 Solidity contract (temel yapı)", "Hafta 1", "✅ Tamamlandı"],
        ["ABI dosyaları üretimi", "Hafta 1", "✅ Tamamlandı"],
        ["Time-lock (24 saat bekleme)", "Hafta 2", "✅ Tamamlandı"],
        ["emergencyPause (acil durdurma)", "Hafta 2", "✅ Tamamlandı"],
        ["Replay koruması (EIP-712 + nonce)", "Hafta 2", "✅ Tamamlandı"],
        ["Slither güvenlik audit", "Hafta 3", "✅ Tamamlandı — 0 Critical, 0 High"],
        ["Deploy script + Hardhat Sepolia config", "Hafta 4 hazırlık", "✅ Hazır"],
        ["Sepolia deploy + Etherscan verify", "Hafta 4 (26 Mayıs)", "⏳ Yarın yapılacak"],
    ],
    header_bg="185FA5",
    col_widths=[9, 4, 5.5],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM 1 — DOSYA YAPISI
# ══════════════════════════════════════════════════════════════════════════════

add_heading("1. Dosya Yapısı ve Mimari", level=1, color=BLUE)

add_para("Deploy edilen contract sayısı: 1 (MultiSigWallet)", bold=True)
add_para(
    "OwnerManager.sol ve TransactionManager.sol ayrı deploy edilmez. "
    "Bunlar 'abstract contract' — MultiSigWallet bu ikisini miras alır ve "
    "tek bir contract olarak Sepolia'ya gönderilir.",
    size=10
)

add_code("MultiSigWallet")
add_code("  ├── is ReentrancyGuard   (OpenZeppelin — reentrancy koruması)")
add_code("  ├── is OwnerManager      (abstract — sahip yönetimi)")
add_code("  └── is TransactionManager (abstract — işlem yaşam döngüsü + time-lock)")

doc.add_paragraph()

add_table(
    headers=["Dosya", "Tür", "Deploy?", "Açıklama"],
    rows=[
        ["MultiSigWallet.sol", "Concrete", "✅ Evet", "Tek deploy edilen contract"],
        ["OwnerManager.sol",   "Abstract", "❌ Hayır", "Sahip ekleme/çıkarma mantığı"],
        ["TransactionManager.sol", "Abstract", "❌ Hayır", "İşlem struct, state, time-lock"],
        ["abi/MultiSigWallet.abi.json", "ABI", "—", "Emine'nin kullandığı ABI (güncel)"],
        ["scripts/deploy.js", "Script", "—", "Sepolia deploy — yarın çalıştırılacak"],
        ["scripts/verify.js", "Script", "—", "Etherscan doğrulama yardımcısı"],
        ["security_audit.md", "Rapor", "—", "Slither audit detayları"],
    ],
    header_bg="185FA5",
    col_widths=[5.5, 3, 3, 7],
)

# EMİNE NOTU
add_heading("Emine'ye — ABI Kullanımı", level=3, color=GREEN)
add_para(
    "Sadece MultiSigWallet ABI'sini kullan: abi/MultiSigWallet.abi.json\n"
    "OwnerManager ve TransactionManager'ın fonksiyonları bu ABI içinde "
    "zaten var — ayrı ABI dosyalarına gerek yok.",
    size=10, color=GREEN
)
add_code("from web3 import Web3")
add_code("with open('abi/MultiSigWallet.abi.json') as f:")
add_code("    abi = json.load(f)")
add_code("contract = w3.eth.contract(address=CONTRACT_ADDRESS, abi=abi)")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM 2 — HAFTA 1: TEMEL KONTRAT YAPISI
# ══════════════════════════════════════════════════════════════════════════════

add_heading("2. Hafta 1 — Temel Kontrat Yapısı", level=1, color=BLUE)

# 2.1 Constructor
add_heading("2.1 Constructor", level=2, color=DARK)
add_para("Cüzdanı başlatan fonksiyon. Deploy sırasında çağrılır.", size=10)
add_code("constructor(address[] memory _owners, uint256 _numConfirmationsRequired)")
doc.add_paragraph()
add_table(
    headers=["Parametre", "Tip", "Açıklama"],
    rows=[
        ["_owners", "address[]", "N adet sahip adresi — tekrarsız, sıfır adres içermez"],
        ["_numConfirmationsRequired", "uint256", "M değeri — gereken minimum onay sayısı"],
    ],
    header_bg="185FA5",
    col_widths=[5.5, 4, 9],
)
add_para("M ≥ N/2 zorunlu kural: örn. 3 sahip → en az 2 onay. Bu kural ihlal edilirse contract deploy olmaz.", bold=True, size=10)

# EMİNE
add_heading("Emine'ye — Deploy Örneği", level=3, color=GREEN)
add_code("owners = [addr1, addr2, addr3]")
add_code("M = 2   # 2-of-3")
add_code("contract = w3.eth.contract(abi=ABI, bytecode=BYTECODE)")
add_code("tx = contract.constructor(owners, M).transact({'from': deployer})")

# 2.2 Core Fonksiyonlar
add_heading("2.2 Core İşlem Fonksiyonları", level=2, color=DARK)
add_table(
    headers=["Fonksiyon", "Kim Çağırır", "Ne Yapar"],
    rows=[
        ["submitTransaction(to, value, data)", "Sahip", "Yeni işlem teklifi oluşturur — txIndex döndürür"],
        ["confirmTransaction(txIndex)",         "Sahip", "İşlemi onaylar; M'e ulaşınca time-lock başlar"],
        ["executeTransaction(txIndex)",          "Sahip", "M onay + time-lock tamam ise işlemi yürütür"],
        ["revokeConfirmation(txIndex)",          "Sahip", "Daha önce verilen onayı geri çeker"],
        ["isConfirmed(txIndex)",                 "Herkes", "İşlem M onaya ulaştı mı? → bool"],
        ["getBalance()",                          "Herkes", "Cüzdan ETH bakiyesi (wei)"],
    ],
    header_bg="185FA5",
    col_widths=[6.5, 3, 9],
)

# 2.3 Sahip Yönetimi
add_heading("2.3 Sahip Yönetimi Fonksiyonları (onlyWallet)", level=2, color=DARK)
add_para(
    "Bu fonksiyonlar doğrudan çağrılamaz. Önce submitTransaction ile teklif "
    "oluşturulur, M onay alınır, executeTransaction ile yürütülür — o sırada "
    "kontrat kendi kendini çağırır (onlyWallet = msg.sender == address(this)).",
    size=10
)
add_table(
    headers=["Fonksiyon", "Açıklama"],
    rows=[
        ["addOwner(address)",                   "Yeni sahip ekler — M ≥ N/2 korunur"],
        ["removeOwner(address)",                 "Sahibi çıkarır — gerekirse M otomatik düşürülür"],
        ["replaceOwner(address, address)",        "Bir sahibi başkasıyla değiştirir"],
        ["changeRequirement(uint256)",            "Onay eşiğini günceller — M ≥ N/2 zorunlu"],
        ["changeTimeLockDuration(uint256)",       "Time-lock süresini değiştirir (saniye)"],
    ],
    header_bg="185FA5",
    col_widths=[6.5, 12],
)

# 2.4 View / Sorgu Fonksiyonları
add_heading("2.4 Sorgu Fonksiyonları", level=2, color=DARK)
add_table(
    headers=["Fonksiyon", "Döndürür", "Açıklama"],
    rows=[
        ["getTransaction(txIndex)", "(address, uint256, bytes, bool, uint256)", "İşlem detayları"],
        ["getTransactionCount()",   "uint256",   "Toplam işlem sayısı"],
        ["getConfirmationStatus(txIndex, owner)", "bool", "O sahip o işlemi onayladı mı?"],
        ["getOwners()",             "address[]", "Tüm sahiplerin listesi"],
        ["getChainId()",            "uint256",   "Mevcut zincir ID"],
        ["getTransactionHash(txIndex)", "bytes32", "EIP-712 uyumlu imzalanacak hash"],
    ],
    header_bg="185FA5",
    col_widths=[5.5, 5.5, 7.5],
)

# 2.5 Events
add_heading("2.5 Events", level=2, color=DARK)
add_table(
    headers=["Event", "Parametreler", "Ne Zaman Yayınlanır"],
    rows=[
        ["Deposit",             "sender, amount, balance",              "ETH yatırılınca"],
        ["SubmitTransaction",   "txIndex, submitter, to, value, data",  "Yeni işlem teklifi oluşunca"],
        ["ConfirmTransaction",  "txIndex, confirmer, confirmationCount","Onay verilince"],
        ["RevokeConfirmation",  "txIndex, revoker",                     "Onay geri çekilince"],
        ["ExecuteTransaction",  "txIndex, executor",                    "İşlem yürütülünce"],
        ["TimeLockTriggered",   "txIndex, unlockTime",                  "M onaya ilk ulaşılınca ★ Yeni"],
        ["TimeLockDurationChanged", "newDuration",                      "Time-lock süresi değişince ★ Yeni"],
        ["EmergencyPaused",     "triggeredBy",                          "Cüzdan dondurulunca ★ Yeni"],
        ["EmergencyUnpaused",   "triggeredBy",                          "Cüzdan açılınca ★ Yeni"],
        ["OwnerAddition",       "owner",                                "Sahip eklenince"],
        ["OwnerRemoval",        "owner",                                "Sahip çıkarılınca / değiştirilince"],
        ["RequirementChange",   "required",                             "M eşiği değişince"],
    ],
    header_bg="185FA5",
    col_widths=[5, 6, 7.5],
)

# EMİNE — Event listener
add_heading("Emine'ye — Event Dinleme", level=3, color=GREEN)
add_para("Tüm eventleri Web3.py ile dinleyebilirsin:", size=10, color=GREEN)
add_code("# Yeni işlem gelince")
add_code("filt = contract.events.SubmitTransaction.create_filter(fromBlock='latest')")
add_code("")
add_code("# M onaya ulaştı, time-lock başladı → UI sayaç göster")
add_code("filt2 = contract.events.TimeLockTriggered.create_filter(fromBlock='latest')")
add_code("# event.args['unlockTime'] → ne zaman execute edilebilir")
add_code("")
add_code("# Cüzdan donduruldu → UI'da uyarı göster")
add_code("filt3 = contract.events.EmergencyPaused.create_filter(fromBlock='latest')")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM 3 — HAFTA 2: GÜVENLİK MEKANİZMALARI
# ══════════════════════════════════════════════════════════════════════════════

add_heading("3. Hafta 2 — Güvenlik Mekanizmaları", level=1, color=BLUE)

# 3.1 Time-lock
add_heading("3.1 Time-lock (24 Saat Bekleme)", level=2, color=DARK)
add_para(
    "Bir işlem M onaya ulaşınca 24 saatlik bir sayaç başlar. "
    "Bu 24 saat dolmadan executeTransaction çalışmaz ve revert atar.",
    size=10
)
add_para("Neden gerekli?", bold=True, size=10)
add_bullet("Onaylar toplansa bile acele hareket etmeyi engeller.")
add_bullet("Fark edilmeyen kötü amaçlı bir işlem için müdahale süresi tanır.")
add_bullet("Büyük tutarlı işlemlerde ek güvence sağlar.")
doc.add_paragraph()
add_para("Nasıl çalışır?", bold=True, size=10)
add_bullet("confirmTransaction çağrılır → onay sayısı M'e ulaşırsa _setTimeLock() tetiklenir")
add_bullet("readyTime[txIndex] = block.timestamp + 24 saat olarak set edilir")
add_bullet("TimeLockTriggered event yayınlanır")
add_bullet("executeTransaction → timeLockPassed modifier kontrolü yapar → block.timestamp < readyTime ise revert")
doc.add_paragraph()
add_para("Önemli notlar:", bold=True, size=10)
add_bullet("Bir onay geri çekilip yeniden verilse de sayaç sıfırlanmaz — ilk M onaya ulaşıldığı anki timer geçerlidir.")
add_bullet("changeTimeLockDuration(saniye) ile süre değiştirilebilir — ama bu multisig konsensüsü gerektirir.")
add_bullet("timeLockDuration şu an 86400 saniye (24 saat). Sıfır yapmak önerilmez.")
doc.add_paragraph()

add_para("Yeni state değişkenleri:", bold=True, size=10)
add_table(
    headers=["Değişken", "Tip", "Erişim", "Açıklama"],
    rows=[
        ["timeLockDuration", "uint256", "public", "Bekleme süresi (saniye) — varsayılan 86400"],
        ["readyTime(txIndex)", "mapping", "public", "İşlem execute edilebilir olacağı timestamp"],
    ],
    header_bg="185FA5",
    col_widths=[4.5, 3, 3, 8],
)

# EMİNE — Time-lock
add_heading("Emine'ye — Time-lock Entegrasyonu", level=3, color=GREEN)
add_para("TimeLockTriggered event'inden unlockTime'ı al, UI'a aktar:", size=10, color=GREEN)
add_code("# TimeLockTriggered event'inden:")
add_code("unlock_time = event.args['unlockTime']  # unix timestamp")
add_code("import time")
add_code("remaining_seconds = unlock_time - time.time()")
add_code("# remaining_seconds > 0 ise henüz execute edilemez")
add_code("")
add_code("# Mevcut readyTime'ı okumak için:")
add_code("ready = contract.functions.readyTime(tx_index).call()")

# İREM — Time-lock
add_heading("İrem'e — Time-lock Test Senaryoları", level=3, color=ORANGE)
add_para("Test etmen gereken durumlar:", size=10, color=ORANGE)
add_bullet('M onay alınır, hemen executeTransaction çağrılır → "TransactionManager: time-lock suresi dolmadi" revert beklenir', color=ORANGE)
add_bullet("TimeLockTriggered event'i doğru txIndex ve unlockTime ile emit edildi mi?", color=ORANGE)
add_bullet("Test ortamında süreyi hızlandırmak için: evm_increaseTime(86401) + evm_mine kullan", color=ORANGE)

divider()

# 3.2 emergencyPause
add_heading("3.2 emergencyPause (Acil Durdurma)", level=2, color=DARK)
add_para(
    "Herhangi bir sahip cüzdanı tek başına dondurabilir. "
    "Dondurulunca submit, confirm, execute, revoke — hiçbiri çalışmaz. "
    "Herhangi bir sahip yeniden açabilir.",
    size=10
)
add_para("Neden unpause da tek sahiple açılıyor?", bold=True, size=10)
add_para(
    "Teknik zorunluluk: cüzdan dondurulunca executeTransaction çalışmaz. "
    "Eğer unpause için multisig konsensüsü gerekseydi, açmak için execute "
    "gerekir ama execute çalışmıyor — kilitlenme (deadlock) olurdu. "
    "Bu nedenle bilinçli olarak tek sahip kararıyla açılabiliyor.",
    size=10, indent=0.5
)
doc.add_paragraph()
add_table(
    headers=["Fonksiyon", "Kim Çağırır", "Açıklama"],
    rows=[
        ["pause()",   "Herhangi bir sahip", "Cüzdanı dondurur — EmergencyPaused event"],
        ["unpause()", "Herhangi bir sahip", "Duraklatmayı kaldırır — EmergencyUnpaused event"],
        ["paused()",  "Herkes (view)",      "Şu an dondurulmuş mu? → bool"],
    ],
    header_bg="185FA5",
    col_widths=[4, 5, 9.5],
)

# İREM — Pause
add_heading("İrem'e — Pause Test Senaryoları", level=3, color=ORANGE)
add_table(
    headers=["Test Senaryosu", "Beklenen Revert"],
    rows=[
        ["Paused iken submitTransaction", '"MultiSigWallet: cuzcdan durduruldu"'],
        ["Paused iken confirmTransaction", '"MultiSigWallet: cuzcdan durduruldu"'],
        ["Paused iken executeTransaction", '"MultiSigWallet: cuzcdan durduruldu"'],
        ["Paused iken revokeConfirmation", '"MultiSigWallet: cuzcdan durduruldu"'],
        ["Zaten paused iken pause()", '"MultiSigWallet: zaten durduruldu"'],
        ["Paused değilken unpause()", '"MultiSigWallet: zaten aktif"'],
        ["Sahip olmayan pause() çağırır", '"MultiSigWallet: sadece sahip cagirir"'],
    ],
    header_bg="854F0B",
    col_widths=[8.5, 10],
)

divider()

# 3.3 Replay Koruması
add_heading("3.3 Replay Koruması (EIP-712 + Chain ID)", level=2, color=DARK)
add_para("İki tür replay saldırısına karşı koruma:", bold=True, size=10)
add_bullet("Zincirler arası: Sepolia'daki işlemi Mainnet'te tekrarlama")
add_bullet("İmza seti replay: Emine'nin topladığı off-chain imzaları ikinci kez kullanma")
doc.add_paragraph()

add_para("Eklenen korumalar:", bold=True, size=10)
add_table(
    headers=["Özellik", "Açıklama", "Nerede Kullanılır"],
    rows=[
        ["DOMAIN_SEPARATOR", "chainId + address(this) içerir. Her contract & zincirde farklı.", "Emine'nin EIP-712 imzalarında"],
        ["nonce", "Her execute'da 1 artar. Eski imza seti geçersiz kalır.", "executeTransaction içinde"],
        ["DEPLOY_CHAIN_ID", "Deploy sırasındaki chainId saklanır.", "executeTransaction'da block.chainid kontrolü"],
        ["getTransactionHash(txIndex)", "EIP-712 uyumlu mesaj hash'i döndürür.", "Off-chain imza sürecinde"],
    ],
    header_bg="185FA5",
    col_widths=[4.5, 8, 6],
)

# EMİNE — EIP-712
add_heading("Emine'ye — EIP-712 Off-chain İmza Entegrasyonu", level=3, color=GREEN)
add_para(
    "Off-chain imza toplarken şu adımları izle. "
    "nonce her execute'da değiştiği için imzayı execute ÖNCE toplamalısın.",
    size=10, color=GREEN
)
add_code("# 1. Kontraktan güncel nonce'u oku")
add_code("current_nonce = contract.functions.nonce().call()")
add_code("")
add_code("# 2. İmzalanacak hash'i kontraktan al")
add_code("tx_hash = contract.functions.getTransactionHash(tx_index).call()")
add_code("")
add_code("# 3. Domain separator — EIP-712 doğrulaması için")
add_code("domain_sep = contract.functions.DOMAIN_SEPARATOR().call()")
add_code("")
add_code("# 4. Her sahibe bu tx_hash'i imzalat")
add_code("from eth_account import Account")
add_code("signed = Account.signHash(tx_hash, private_key=owner_private_key)")
add_code("")
add_code("# 5. İmzaları topla, sonra confirmTransaction on-chain çağır")
add_code("contract.functions.confirmTransaction(tx_index).transact({'from': owner})")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM 4 — HAFTA 3: SLITHER GÜVENLİK AUDIT
# ══════════════════════════════════════════════════════════════════════════════

add_heading("4. Hafta 3 — Slither Güvenlik Audit", level=1, color=BLUE)
add_para(
    "Slither, Solidity kontratlarını otomatik analiz eden endüstri standardı "
    "bir güvenlik aracıdır. pip install slither-analyzer ile kuruldu ve "
    "tüm kontratlar üzerinde çalıştırıldı.",
    size=10
)
doc.add_paragraph()

add_para("Sonuç özeti:", bold=True, size=10)
add_table(
    headers=["Önem Seviyesi", "İlk Çalıştırma", "Düzeltme Sonrası"],
    rows=[
        ["Critical", "0", "0"],
        ["High",     "0", "0"],
        ["Medium",   "2", "0 ✅ Düzeltildi"],
        ["Low / Info", "5", "5 (false positive / kasıtlı)"],
    ],
    header_bg="185FA5",
    col_widths=[6, 5, 7.5],
)
add_para("Sonuç: 0 Critical, 0 High bulgu. Audit temiz.", bold=True, size=10, color=DARK_GREEN)
doc.add_paragraph()

add_heading("Düzeltilen Sorunlar", level=2, color=DARK)

add_para("[DÜZELTİLDİ] Strict Equality — Medium", bold=True, size=10)
add_para(
    "readyTime[txIndex] == 0 ifadesi riskli bulundu. Düzeltme: "
    "ayrı bir bool _timeLockInitialized mapping eklendi, == 0 kontrolü kaldırıldı.",
    size=10, indent=0.5
)
add_code("// Önce (riskli):")
add_code("if (readyTime[_txIndex] == 0) { ... }")
add_code("")
add_code("// Sonra (güvenli):")
add_code("if (!_timeLockInitialized[_txIndex]) {")
add_code("    _timeLockInitialized[_txIndex] = true;")
add_code("    ...")
add_code("}")

doc.add_paragraph()
add_para("[DÜZELTİLDİ] Array Length Cache — Low (Gas)", bold=True, size=10)
add_para(
    "OwnerManager.sol'daki removeOwner ve replaceOwner döngüleri her "
    "adımda owners.length'i storage'dan okuyordu. Düzeltme: döngüden önce "
    "uint256 len = owners.length ile cache'lendi. Gas tasarrufu sağlandı.",
    size=10, indent=0.5
)

doc.add_paragraph()
add_heading("Kalan Bulgular — False Positive / Kasıtlı", level=2, color=DARK)
add_table(
    headers=["Bulgu", "Neden Düzeltilmedi"],
    rows=[
        ["assembly — OpenZeppelin", "OZ kütüphanesinin kendi kodu, biz yazmadık"],
        ["pragma version", "OZ ^0.8.20 kullanıyor, değiştirilemez"],
        ["solc-version", "Bahsedilen 3 bug projemizi etkilemiyor (kontrol edildi)"],
        ["low-level-calls", "Multisig cüzdan için kaçınılmaz — hedef kontrat compile-time bilinmez"],
        ["naming-convention", "_parametre prefix'i Solidity standardı, CAPS sabitler için de öyle"],
    ],
    header_bg="185FA5",
    col_widths=[5, 13.5],
)

add_para("Detaylı analiz için: security_audit.md dosyasına bakın.", size=10, color=MUTED, italic=True)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM 5 — HAFTA 4: SEPOLIA DEPLOY
# ══════════════════════════════════════════════════════════════════════════════

add_heading("5. Hafta 4 — Sepolia Deploy (26 Mayıs — Yarın)", level=1, color=BLUE)
add_para(
    "Deploy script ve tüm teknik altyapı hazır. Yarın ortak oturumda "
    "yapılacak. Merve deploy eder ve Etherscan'de verify eder.",
    size=10
)
doc.add_paragraph()

add_heading("Emine ve İrem'den Bugün Alınması Gereken Tek Şey:", level=2, color=DARK)
add_para(
    "Her ikinizin MetaMask cüzdan adresiniz (0x ile başlayan 42 karakter). "
    "Bu adresler MultiSigWallet'ın owner listesine girecek.",
    size=10, bold=True
)
doc.add_paragraph()

add_heading("Deploy Adımları", level=2, color=DARK)
add_para("1. Aşağıdaki servislere ücretsiz hesap aç (sadece Merve):", bold=True, size=10)
add_bullet("Infura (infura.io) veya Alchemy (alchemy.com) → Sepolia RPC URL")
add_bullet("Etherscan (etherscan.io) → API Key")
add_bullet("sepoliafaucet.com → Test ETH al")
doc.add_paragraph()
add_para("2. .env dosyasını oluştur (.env.example'ı kopyala, doldur):", bold=True, size=10)
add_code("SEPOLIA_RPC_URL   = Infura/Alchemy URL")
add_code("PRIVATE_KEY       = Merve'nin MetaMask private key'i")
add_code("OWNER_1           = Merve'nin adresi")
add_code("OWNER_2           = Emine'nin adresi  ← yarın alınacak")
add_code("OWNER_3           = İrem'in adresi    ← yarın alınacak")
add_code("ETHERSCAN_API_KEY = Etherscan API key")
doc.add_paragraph()
add_para("3. Deploy:", bold=True, size=10)
add_code("node scripts/deploy.js")
add_para("→ Contract adresi terminalde görünür ve scripts/deploy_info.json'a kaydedilir.", size=10, indent=0.5)
doc.add_paragraph()
add_para("4. Verify:", bold=True, size=10)
add_code("node scripts/verify.js")
add_para("→ Script Etherscan'e doğrulama için adımları gösterir.", size=10, indent=0.5)
doc.add_paragraph()
add_para("5. Herkes için demo:", bold=True, size=10)
add_bullet("Merve (Owner 1): işlem önerir")
add_bullet("Emine (Owner 2): backend üzerinden onaylar")
add_bullet("İrem (Owner 3): Streamlit UI üzerinden onaylar")

# EMİNE — deploy sonrası
add_heading("Emine'ye — Deploy Sonrası Yapman Gerekenler", level=3, color=GREEN)
add_para("Merve sana CONTRACT_ADDRESS'i verecek. Bunu .env'ye ekle:", size=10, color=GREEN)
add_code("CONTRACT_ADDRESS=0x...  # Merve'nin deploy ettiği adres")
add_code("CHAIN_ID=11155111       # Sepolia")
add_para("web3.py bağlantısını Sepolia'ya güncelle:", size=10, color=GREEN)
add_code("w3 = Web3(Web3.HTTPProvider(os.getenv('SEPOLIA_RPC_URL')))")
add_code("assert w3.eth.chain_id == 11155111, 'Sepolia değil!'")

# İREM — deploy sonrası
add_heading("İrem'e — Deploy Sonrası Yapman Gerekenler", level=3, color=ORANGE)
add_para(
    "Contract adresi belli olunca Streamlit UI'ındaki sabit adresi "
    "güncelle. Test ağı olduğu için Etherscan Sepolia'da işlemleri "
    "görebilirsin:",
    size=10, color=ORANGE
)
add_code("https://sepolia.etherscan.io/address/CONTRACT_ADDRESS")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM 6 — TAM REVERT MESAJLARI LİSTESİ
# ══════════════════════════════════════════════════════════════════════════════

add_heading("6. Tüm Revert Mesajları (İrem için — Test)", level=1, color=BLUE)
add_para(
    "Aşağıdaki tablo projedeki her revert mesajını listeler. "
    "Test yazarken bu mesajları birebir kullan.",
    size=10
)
doc.add_paragraph()

add_table(
    headers=["Senaryo", "Revert Mesajı", "★"],
    rows=[
        ["Boş owners dizisi",
         "MultiSigWallet: en az 1 sahip gerekli", ""],
        ["M < N/2 veya M=0",
         "MultiSigWallet: gecersiz M veya N degeri (M >= N/2 sarti)", ""],
        ["address(0) owner",
         "MultiSigWallet: sifir adres", ""],
        ["Tekrar eden owner",
         "MultiSigWallet: sahip adresi tekrarlanamaz", ""],
        ["Sahip olmayan çağırır",
         "MultiSigWallet: sadece sahip cagirir", ""],
        ["Sıfır hedef adres (submit)",
         "MultiSigWallet: sifir hedef adres", ""],
        ["Olmayan txIndex",
         "TransactionManager: islem bulunamadi", ""],
        ["Zaten yürütülmüş tx",
         "TransactionManager: islem zaten yurutuldu", ""],
        ["Aynı sahip 2 kez onaylar",
         "TransactionManager: zaten onaylanmis", ""],
        ["Onayı olmayan revoke",
         "TransactionManager: onay bulunamadi", ""],
        ["Yetersiz onay ile execute",
         "MultiSigWallet: yetersiz onay sayisi", ""],
        ["addOwner direkt çağrılır",
         "OwnerManager: sadece kontrat cagirabilir", ""],
        ["Time-lock dolmadan execute",
         "TransactionManager: time-lock suresi dolmadi", "★ Yeni"],
        ["Paused iken işlem (4 fonksiyon)",
         "MultiSigWallet: cuzcdan durduruldu", "★ Yeni"],
        ["Zaten paused iken pause()",
         "MultiSigWallet: zaten durduruldu", "★ Yeni"],
        ["Paused değilken unpause()",
         "MultiSigWallet: zaten aktif", "★ Yeni"],
        ["Yanlış zincirde execute",
         "MultiSigWallet: yanlis zincir", "★ Yeni"],
        ["Execute başarısız (hedef revert)",
         "MultiSigWallet: islem yurutme basarisiz", ""],
    ],
    header_bg="854F0B",
    col_widths=[6, 10, 2.5],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM 7 — EMİNE İÇİN ÖZET
# ══════════════════════════════════════════════════════════════════════════════

add_heading("7. Emine İçin Özet — Backend Entegrasyonu", level=1, color=GREEN)

add_para("Kullanacağın tek ABI:", bold=True, size=10)
add_code("abi/MultiSigWallet.abi.json")

doc.add_paragraph()
add_para("Deploy parametreleri:", bold=True, size=10)
add_code("owners = [addr_merve, addr_emine, addr_irem]")
add_code("M = 2  # 2-of-3")

doc.add_paragraph()
add_para("Yeni eklenen — kullanman gereken şeyler:", bold=True, size=10, color=GREEN)
add_table(
    headers=["Özellik", "Nasıl Kullanırsın"],
    rows=[
        ["paused()", "İşlem göndermeden önce kontrol et — True ise UI'da uyarı göster"],
        ["pause() / unpause()", "Acil durum butonları için"],
        ["readyTime(txIndex)", "İşlemin ne zaman execute edilebileceğini hesapla"],
        ["TimeLockTriggered event", "UI'da 'X saat sonra yürütülebilir' göstergesi için"],
        ["nonce()", "EIP-712 off-chain imza toplarken güncel nonce'u al"],
        ["getTransactionHash(txIndex)", "İmzalanacak hash'i kontraktan al"],
        ["DOMAIN_SEPARATOR()", "EIP-712 domain separator — off-chain imzada kullan"],
        ["getChainId()", "Doğru zincirde olduğunu kontrol et"],
        ["EmergencyPaused / Unpaused event", "UI'da cüzdan durumu göstergesi için"],
    ],
    header_bg="0F6E56",
    col_widths=[5.5, 13],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# BÖLÜM 8 — İREM İÇİN ÖZET
# ══════════════════════════════════════════════════════════════════════════════

add_heading("8. İrem İçin Özet — Test & UI", level=1, color=ORANGE)

add_para("Test yazarken bilmen gerekenler:", bold=True, size=10, color=ORANGE)
doc.add_paragraph()

add_table(
    headers=["Test Kategorisi", "Test Senaryosu", "Beklenen"],
    rows=[
        ["Constructor", "M < N/2 ile deploy", "Revert"],
        ["Constructor", "Tekrar eden owner", "Revert"],
        ["Time-lock", "M onay sonrası hemen execute", "Revert: time-lock"],
        ["Time-lock", "evm_increaseTime(86401) sonrası execute", "Başarılı"],
        ["Time-lock", "TimeLockTriggered event doğrulama", "Event doğru emit"],
        ["Pause", "Paused iken submit/confirm/execute/revoke", "Revert: cuzcdan durduruldu"],
        ["Pause", "Sahip olmayan pause()", "Revert: sadece sahip cagirir"],
        ["Replay", "Aynı txIndex'i iki kez execute", "Revert: zaten yurutuldu"],
        ["Replay", "nonce execute sonrası 1 arttı mı?", "nonce == önceki + 1"],
        ["Reentrancy", "Kötü amaçlı kontrat geri çağırır", "Revert (ReentrancyGuard)"],
        ["2-of-3 flow", "Submit → 2 onay → execute (tam akış)", "Başarılı"],
        ["3-of-5 flow", "Submit → 2 onay → execute (yetersiz)", "Revert"],
    ],
    header_bg="854F0B",
    col_widths=[3.5, 7.5, 7.5],
)

doc.add_paragraph()
add_para("UI için kullanman gereken fonksiyonlar:", bold=True, size=10, color=ORANGE)
add_table(
    headers=["UI Bileşeni", "Contract Çağrısı"],
    rows=[
        ["Pending tx listesi", "getTransactionCount() + getTransaction(i) ile loop"],
        ["Onay göstergesi (kim onayladı)", "getConfirmationStatus(txIndex, ownerAddr)"],
        ["Cüzdan bakiyesi", "getBalance()"],
        ["Time-lock sayaç", "readyTime(txIndex) — unix timestamp"],
        ["Dondurulmuş mu göstergesi", "paused()"],
        ["Sahip listesi", "getOwners()"],
        ["Canlı güncelleme", "SubmitTransaction, ConfirmTransaction, TimeLockTriggered, EmergencyPaused eventleri"],
    ],
    header_bg="854F0B",
    col_widths=[5, 13.5],
)

# ── Kaydet ───────────────────────────────────────────────────────────────────

out_path = r"c:\Users\merve\Desktop\8. dönem\blockchain\blockchain-prohe\Merve_Ilerleme_Raporu.docx"
doc.save(out_path)
print(f"Rapor kaydedildi: {out_path}")
