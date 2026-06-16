"""
Rapor Bölüm 1-3 + Güvenlik Analizi DOCX Oluşturucu
COM 4532 Blockchain Projesi — Merve Çakır
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Sayfa kenar boşlukları ────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Renk sabitleri ────────────────────────────────────────
DARK_BLUE  = RGBColor(0x18, 0x5F, 0xA5)
DARK_GREEN = RGBColor(0x0F, 0x6E, 0x56)
DARK_GRAY  = RGBColor(0x33, 0x33, 0x33)
MID_GRAY   = RGBColor(0x66, 0x66, 0x66)
LIGHT_BLUE = RGBColor(0xE6, 0xF1, 0xFB)
RED        = RGBColor(0xC0, 0x00, 0x00)
BLACK      = RGBColor(0x00, 0x00, 0x00)

# ── Yardımcı fonksiyonlar ─────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=DARK_BLUE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold      = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)
    return p

def add_body(doc, text, bold=False, italic=False, color=DARK_GRAY, size=10.5):
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    return p

def add_bullet(doc, text, color=DARK_GRAY):
    p   = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    return p

def add_code(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name       = 'Courier New'
    run.font.size       = Pt(9)
    run.font.color.rgb  = RGBColor(0x1E, 0x40, 0x6E)
    p.paragraph_format.left_indent   = Cm(1)
    p.paragraph_format.space_after   = Pt(2)
    p.paragraph_format.space_before  = Pt(2)

def add_ref(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"[{num}] ")
    r1.font.size      = Pt(9.5)
    r1.font.bold      = True
    r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(text)
    r2.font.size      = Pt(9.5)
    r2.font.color.rgb = MID_GRAY

# ════════════════════════════════════════════════════════════
# KAPAK
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("COM 4532 — Blockchain Teknolojileri")
r.font.size = Pt(12); r.font.color.rgb = MID_GRAY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Multi-Signature Wallet")
r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Teknik Rapor — Bölüm 1–3: Giriş, Mimari ve Güvenlik Analizi")
r.font.size = Pt(13); r.font.color.rgb = DARK_GREEN; r.font.italic = True

doc.add_paragraph()

tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
bilgiler = [
    ("Hazırlayan",   "Merve Çakır"),
    ("Ders",         "COM 4532 — Blockchain Teknolojileri"),
    ("Proje",        "Multi-Signature Wallet (2-of-3)"),
    ("Testnet",      "Ethereum Sepolia"),
    ("Kontrat",      "0x883f6b9ea658D8c74999C788903677fB7eAE9a1e"),
]
for i, (k, v) in enumerate(bilgiler):
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    tbl.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = DARK_BLUE
    set_cell_bg(tbl.rows[i].cells[0], 'E6F1FB')
for row in tbl.rows:
    for cell in row.cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# BÖLÜM 1 — GİRİŞ VE PROBLEM TANIMI
# ════════════════════════════════════════════════════════════
add_heading(doc, "1. Giriş ve Problem Tanımı", level=1)

add_body(doc,
    "Dijital varlıkların güvenli yönetimi, blockchain teknolojisinin temel problemlerinden birini oluşturmaktadır. "
    "Geleneksel tek imzalı (single-signature) cüzdanlarda, özel anahtar (private key) sahibi tüm kontrole sahiptir; "
    "bu anahtar çalındığında veya kaybolduğunda varlıklar kalıcı olarak erişilemez hale gelir [1]. "
    "2014 yılında Mt. Gox borsasının 850.000 Bitcoin kaybetmesi ve 2016 yılındaki DAO hack'i (60 milyon USD), "
    "tek yetkili kontrol modelinin getirdiği güvenlik açıklarının somut örnekleridir [2, 3].")

add_heading(doc, "1.1 Problem Tanımı", level=2)

add_body(doc,
    "Merkezi yönetim yapılarında gözlemlenen başlıca güvenlik riskleri şu şekilde özetlenebilir:")
for item in [
    "Tek nokta başarısızlığı (Single Point of Failure): Tek bir özel anahtarın ele geçirilmesi tüm varlıkların kaybına yol açar.",
    "İçeriden tehdit (Insider Threat): Yetkili kişinin kötü niyetli davranışına karşı denetim mekanizması yoktur.",
    "Operasyonel hata: Yanlış adrese gönderim gibi insan hatalarını geri almak mümkün değildir.",
    "Anahtar kaybı: Özel anahtarın kaybolması durumunda varlıklara erişim kalıcı olarak kesilir.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.2 Önerilen Çözüm: Multi-Signature Wallet", level=2)

add_body(doc,
    "Bu projede Ethereum blockchain üzerinde çalışan bir Çok İmzalı Cüzdan (Multi-Signature Wallet) tasarlanmış "
    "ve Solidity 0.8.20 ile implement edilmiştir. Sistem, M-of-N eşiği modeline dayanmaktadır: N adet yetkili "
    "sahipten en az M tanesinin onayı olmadan hiçbir işlem gerçekleştirilemez [4].")

add_body(doc,
    "Projede 2-of-3 konfigürasyonu benimsenmiştir: üç sahipten en az ikisinin onayı zorunludur. "
    "Bu yapı hem güvenliği hem de operasyonel sürekliliği sağlar — bir sahip geçici olarak erişilemez olsa "
    "bile diğer iki sahip sistemi işletmeye devam edebilir.")

add_heading(doc, "1.3 Katkılar ve Yenilikler", level=2)

add_body(doc, "Bu proje, temel multi-sig mimarisinin ötesinde aşağıdaki güvenlik mekanizmalarını da içermektedir:")
for item in [
    "24 saatlik zaman kilidi (time-lock): M onaya ulaşıldıktan sonra işlem belirlenen süre dolmadan yürütülemez.",
    "Acil durdurma (emergencyPause): Güvenlik tehdidi algılandığında herhangi bir sahip sistemi anında dondurabilir.",
    "EIP-712 replay koruması: Zincirler arası ve tekrarlı saldırılara karşı domain separator ve nonce mekanizması.",
    "Zincir kimliği doğrulama (chainId check): İşlemler yalnızca deploy edildiği zincirde geçerlidir.",
]:
    add_bullet(doc, item)

add_heading(doc, "1.4 Teknoloji Yığını", level=2)

tbl = doc.add_table(rows=7, cols=2)
tbl.style = 'Table Grid'
satirlar = [
    ("Bileşen", "Teknoloji / Sürüm"),
    ("Akıllı Kontrat Dili", "Solidity 0.8.20"),
    ("Geliştirme Ortamı", "Hardhat 3.x (ESM modülü)"),
    ("Güvenlik Kütüphanesi", "OpenZeppelin Contracts v5"),
    ("Test Ağı", "Ethereum Sepolia Testnet (chainId: 11155111)"),
    ("Backend Entegrasyonu", "Web3.py 7.16.0 (Python)"),
    ("Güvenlik Analizi", "Slither v0.11.5"),
]
for i, (k, v) in enumerate(satirlar):
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    if i == 0:
        set_cell_bg(tbl.rows[i].cells[0], '185FA5')
        set_cell_bg(tbl.rows[i].cells[1], '185FA5')
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tbl.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        tbl.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
    else:
        set_cell_bg(tbl.rows[i].cells[0], 'E6F1FB')
    for cell in tbl.rows[i].cells:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# BÖLÜM 2 — MİMARİ AÇIKLAMA
# ════════════════════════════════════════════════════════════
add_heading(doc, "2. Sistem Mimarisi", level=1)

add_body(doc,
    "Sistem mimarisi üç katmandan oluşmaktadır: Solidity akıllı kontrat katmanı, Python/Web3.py backend katmanı "
    "ve Streamlit kullanıcı arayüzü katmanı. Bu bölümde ağırlıklı olarak akıllı kontrat mimarisi ele alınmaktadır.")

add_heading(doc, "2.1 Akıllı Kontrat Yapısı", level=2)

add_body(doc,
    "Kontrat mimarisi nesne yönelimli tasarım prensipleri çerçevesinde üç ayrı soyut kontrata ayrılmıştır. "
    "Bu ayrım, tek sorumluluk ilkesini (Single Responsibility Principle) Solidity'e uyarlamaktadır [5]:")

add_body(doc, "Kontrat Hiyerarşisi:", bold=True)
add_code(doc, "OwnerManager (abstract)")
add_code(doc, "    └── TransactionManager (abstract)")
add_code(doc, "            └── MultiSigWallet  ← deploy edilen kontrat")

add_heading(doc, "2.1.1 OwnerManager.sol", level=3)
add_body(doc,
    "Sahip (owner) yaşam döngüsünü yönetir. addOwner(), removeOwner(), replaceOwner() ve "
    "changeRequirement() fonksiyonlarını içerir. Tüm bu fonksiyonlar onlyWallet modifier'ı ile korunmaktadır; "
    "yani yalnızca kontratın kendisi (çok imza süreci tamamlandıktan sonra) bu fonksiyonları çağırabilir.")

for item in [
    "isOwner mapping: Adresin sahip olup olmadığını O(1) sürede kontrol eder.",
    "owners dizisi: Tüm sahipleri sıralı şekilde tutar (getOwners() için).",
    "numConfirmationsRequired: M değeri — kaç onay gerektiğini saklar.",
    "M ≥ N/2 kuralı: Constructor'da ve her değişiklikte doğrulanır.",
]:
    add_bullet(doc, item)

add_heading(doc, "2.1.2 TransactionManager.sol", level=3)
add_body(doc,
    "İşlem (transaction) yaşam döngüsünü yönetir. Her işlem bir Transaction struct'ında saklanır "
    "ve transactions[] dizisine eklenir.")

for item in [
    "Transaction struct: to, value, data, executed, numConfirmations alanlarını içerir.",
    "confirmed mapping: Her işlem için hangi sahibin onay verdiğini takip eder.",
    "readyTime mapping: Time-lock için işlemin en erken yürütülebileceği zamanı saklar.",
    "_timeLockInitialized mapping: Slither bulgusuna göre eklenen boolean flag.",
    "timeLockDuration: Varsayılan 24 saat, onlyWallet ile değiştirilebilir.",
]:
    add_bullet(doc, item)

add_heading(doc, "2.1.3 MultiSigWallet.sol", level=3)
add_body(doc,
    "Ana kontrat — OwnerManager ve TransactionManager'ı miras alır. "
    "OpenZeppelin'in ReentrancyGuard'ını da uygular [6]. Aşağıdaki kritik state değişkenlerini ekler:")

for item in [
    "paused (bool): emergencyPause durumunu takip eder.",
    "nonce (uint256): Her execute'da artar; replay saldırısını önler.",
    "DOMAIN_SEPARATOR (bytes32 immutable): EIP-712 standartına uygun, constructor'da hesaplanır [7].",
    "DEPLOY_CHAIN_ID (uint256 immutable): Deploy anında kaydedilir; zincirler arası replay'e karşı.",
]:
    add_bullet(doc, item)

add_heading(doc, "2.2 İşlem Akışı", level=2)

add_body(doc, "Bir işlemin tamamlanması beş adımdan oluşur:")

adimlar = [
    ("1. submitTransaction()", "Herhangi bir sahip hedef adres, ETH miktarı ve opsiyonel calldata ile teklif açar."),
    ("2. confirmTransaction()", "Sahipler teker teker onay verir. M. onaya ulaşıldığında time-lock başlar."),
    ("3. Time-lock bekleme", "Belirlenen süre (varsayılan 24 saat) dolana kadar işlem yürütülemez."),
    ("4. executeTransaction()", "Süre dolduktan sonra herhangi bir sahip işlemi yürütür; CEI pattern uygulanır."),
    ("5. Event yayını", "ExecuteTransaction eventi yayınlanır; backend ve UI bu eventi dinleyerek güncellenir."),
]
tbl = doc.add_table(rows=len(adimlar)+1, cols=2)
tbl.style = 'Table Grid'
tbl.rows[0].cells[0].text = "Adım"
tbl.rows[0].cells[1].text = "Açıklama"
set_cell_bg(tbl.rows[0].cells[0], '0F6E56')
set_cell_bg(tbl.rows[0].cells[1], '0F6E56')
for cell in tbl.rows[0].cells:
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
for i, (adim, acik) in enumerate(adimlar):
    tbl.rows[i+1].cells[0].text = adim
    tbl.rows[i+1].cells[1].text = acik
    tbl.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
    tbl.rows[i+1].cells[0].paragraphs[0].runs[0].font.color.rgb = DARK_GREEN
    set_cell_bg(tbl.rows[i+1].cells[0], 'E1F5EE')
    for cell in tbl.rows[i+1].cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, "2.3 Güvenlik Tasarım Desenleri", level=2)

add_heading(doc, "2.3.1 CEI (Checks-Effects-Interactions) Pattern", level=3)
add_body(doc,
    "executeTransaction() fonksiyonunda CEI deseni titizlikle uygulanmıştır [8]. "
    "Bu desen, reentrancy saldırısının temel savunma mekanizmasıdır:")
add_code(doc, "// CHECKS")
add_code(doc, "require(!txn.executed, 'zaten yurutuldu');")
add_code(doc, "require(block.chainid == DEPLOY_CHAIN_ID, 'yanlis zincir');")
add_code(doc, "// EFFECTS (dış çağrıdan ÖNCE state güncellenir)")
add_code(doc, "txn.executed = true;")
add_code(doc, "nonce++;")
add_code(doc, "// INTERACTIONS (en sonda dış çağrı)")
add_code(doc, "(bool success,) = txn.to.call{value: txn.value}(txn.data);")

add_heading(doc, "2.3.2 EIP-712 Domain Separator", level=3)
add_body(doc,
    "Off-chain imzalama için EIP-712 standartı kullanılmıştır. Domain separator, kontrat adını, "
    "versiyonunu, zincir ID'sini ve kontrat adresini içerir; bu sayede imza başka bir kontrat veya "
    "zincirde geçersizdir [7]:")
add_code(doc, "DOMAIN_SEPARATOR = keccak256(abi.encode(")
add_code(doc, "    keccak256('EIP712Domain(string name,string version,")
add_code(doc, "               uint256 chainId,address verifyingContract)'),")
add_code(doc, "    keccak256('MultiSigWallet'), keccak256('1'),")
add_code(doc, "    block.chainid, address(this)")
add_code(doc, "));")

add_heading(doc, "2.4 Deploy Bilgileri", level=2)

tbl2 = doc.add_table(rows=5, cols=2)
tbl2.style = 'Table Grid'
deploy_bilgi = [
    ("Kontrat Adresi",  "0x883f6b9ea658D8c74999C788903677fB7eAE9a1e"),
    ("Deploy TX Hash",  "0x72e0a4b15333...220236f73"),
    ("Ağ",              "Ethereum Sepolia Testnet (chainId: 11155111)"),
    ("Deploy Tarihi",   "09 Haziran 2026"),
    ("Gas Kullanımı",   "2,252,408 gas"),
]
for i, (k, v) in enumerate(deploy_bilgi):
    tbl2.rows[i].cells[0].text = k
    tbl2.rows[i].cells[1].text = v
    set_cell_bg(tbl2.rows[i].cells[0], 'E6F1FB')
    tbl2.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    tbl2.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = DARK_BLUE
    for cell in tbl2.rows[i].cells:
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# BÖLÜM 3 — GÜVENLİK ANALİZİ
# ════════════════════════════════════════════════════════════
add_heading(doc, "3. Güvenlik Analizi", level=1)

add_body(doc,
    "Güvenlik analizi iki aşamada gerçekleştirilmiştir: (1) Slither statik analiz aracı ile otomatik tarama, "
    "(2) manuel erişim kontrolü denetimi ve saldırı vektörü değerlendirmesi.")

add_heading(doc, "3.1 Slither Statik Analizi", level=2)

add_body(doc,
    "Slither, Solidity kaynak kodunu statik olarak analiz eden ve yaygın güvenlik açıklarını tespit eden "
    "bir araçtır [9]. Analiz 25 Mayıs 2026 tarihinde Slither v0.11.5 ile gerçekleştirilmiştir. "
    "solc-select aracı kullanılarak derleyici sürümü kontrat ile eşleştirilmiş (0.8.20) ve "
    "OpenZeppelin remapping uygulanmıştır.")

add_heading(doc, "3.1.1 Analiz Özeti", level=3)

tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
ozet = [
    ("Seviye", "İlk Çalıştırma", "Düzeltme Sonrası"),
    ("Critical", "0", "0"),
    ("High",     "0", "0"),
    ("Medium",   "2", "0 ✓"),
    ("Low / Info","5", "5 (false positive)"),
]
renkler = ['185FA5', 'FFFFFF', 'FFFFFF', 'FFF2CC', 'FFFFFF']
for i, (a, b, c) in enumerate(ozet):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
    tbl.rows[i].cells[2].text = c
    set_cell_bg(tbl.rows[i].cells[0], renkler[i])
    if i == 0:
        for cell in tbl.rows[i].cells:
            set_cell_bg(cell, '185FA5')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            cell.paragraphs[0].runs[0].font.bold = True
    if i == 3:
        tbl.rows[i].cells[2].paragraphs[0].runs[0].font.color.rgb = DARK_GREEN
        tbl.rows[i].cells[2].paragraphs[0].runs[0].font.bold = True
    for cell in tbl.rows[i].cells:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(10)

add_body(doc, "Sonuç: 0 Critical, 0 High bulgu. 2 Medium bulgu düzeltildi.", bold=True, color=DARK_GREEN)

add_heading(doc, "3.1.2 Düzeltilen Bulgular", level=3)

add_body(doc, "M-01 — Katı Eşitlik Karşılaştırması (incorrect-equality) [DÜZELTILDI]",
         bold=True, color=DARK_BLUE)
add_body(doc,
    "Konum: TransactionManager._setTimeLock() — readyTime[_txIndex] == 0 karşılaştırması. "
    "Risk seviyesi: Medium. Sıfır değeriyle katı eşitlik karşılaştırması beklenmedik davranışlara "
    "yol açabilir; özellikle depolama değişkeni henüz yazılmamışsa varsayılan sıfır değeri "
    "yanlış sonuç üretebilir.")
add_body(doc, "Düzeltme: Ayrı bir _timeLockInitialized boolean flag'i eklendi:", italic=True)
add_code(doc, "// ÖNCE (riskli):")
add_code(doc, "if (readyTime[_txIndex] == 0) { ... }")
add_code(doc, "")
add_code(doc, "// SONRA (güvenli):")
add_code(doc, "if (!_timeLockInitialized[_txIndex]) {")
add_code(doc, "    _timeLockInitialized[_txIndex] = true;")
add_code(doc, "    ...")
add_code(doc, "}")

doc.add_paragraph()
add_body(doc, "M-02 — Zaman Damgası Manipülasyonu (timestamp) [DÜZELTILDI]",
         bold=True, color=DARK_BLUE)
add_body(doc,
    "Konum: TransactionManager._setTimeLock() — block.timestamp kullanımı. "
    "Risk seviyesi: Medium. Ethereum madencileri (veya PoS validatörleri) block.timestamp'i "
    "yaklaşık 15 dakika ileri/geri manipüle edebilir [10]. "
    "M-01 ile aynı satırdan kaynaklandığından M-01 düzeltmesiyle birlikte giderildi. "
    "Ek not: Time-lock süresi 24 saat olduğundan 15 dakikalık manipülasyon pratikte önemsizdir.")

doc.add_paragraph()
add_body(doc, "L-01 — Dizi Uzunluğu Cache'lenmemesi (cache-array-length) [DÜZELTILDI]",
         bold=True, color=DARK_BLUE)
add_body(doc,
    "Konum: OwnerManager.removeOwner() ve OwnerManager.replaceOwner() döngüleri. "
    "Risk seviyesi: Low. Her iterasyonda owners.length storage'dan okunması gereksiz gas tüketimine yol açar.")
add_code(doc, "// ÖNCE: Her iterasyonda storage okuma (pahalı)")
add_code(doc, "for (uint256 i = 0; i < owners.length; i++) { ... }")
add_code(doc, "")
add_code(doc, "// SONRA: Tek seferinde cache'le")
add_code(doc, "uint256 len = owners.length;")
add_code(doc, "for (uint256 i = 0; i < len; i++) { ... }")

add_heading(doc, "3.1.3 False Positive / Kasıtlı Bulgular", level=3)

fp_listesi = [
    ("I-01 — Assembly Kullanımı",
     "OpenZeppelin StorageSlot.sol kütüphanesi kodu. Proje tarafından yazılmamıştır; eylem gerekmez."),
    ("I-02 — Pragma Sürüm Uyumsuzluğu",
     "Proje sabit sürüm (0.8.20) kullanırken OZ esnek sürüm (^0.8.20) kullanıyor. "
     "Sabit sürüm güvenlik açısından daha iyidir."),
    ("I-03 — Solidity Sürüm Uyarısı",
     "0.8.20 sürümündeki bilinen üç hata (VerbatimInvalidDeduplication, FullInliner, "
     "MissingSideEffectsOnSelectorAccess) projede kullanılan pattern'leri etkilememektedir."),
    ("I-04 — Low-Level Call",
     "executeTransaction() içindeki txn.to.call{value}(data) kullanımı tasarım gereğidir. "
     "Hedef kontrat compile-time'da bilinmediğinden yüksek seviyeli çağrı kullanılamaz. "
     "ReentrancyGuard ve CEI pattern ile korunmaktadır."),
    ("I-05 — İsimlendirme Kuralları",
     "Parametre isimleri (_txIndex, _to) ve sabitler (DOMAIN_SEPARATOR) Solidity "
     "topluluğunun yaygın konvansiyonuna uygundur."),
]
tbl = doc.add_table(rows=len(fp_listesi)+1, cols=2)
tbl.style = 'Table Grid'
tbl.rows[0].cells[0].text = "Bulgu"
tbl.rows[0].cells[1].text = "Değerlendirme"
for cell in tbl.rows[0].cells:
    set_cell_bg(cell, '666666')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
for i, (k, v) in enumerate(fp_listesi):
    tbl.rows[i+1].cells[0].text = k
    tbl.rows[i+1].cells[1].text = v
    set_cell_bg(tbl.rows[i+1].cells[0], 'F5F5F5')
    tbl.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
    tbl.rows[i+1].cells[0].paragraphs[0].runs[0].font.size = Pt(9.5)
    tbl.rows[i+1].cells[1].paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_paragraph()
doc.add_page_break()

add_heading(doc, "3.2 Erişim Kontrolü Denetimi", level=2)

add_body(doc,
    "Tüm admin fonksiyonları için erişim kontrolü manuel olarak denetlenmiştir. "
    "Aşağıdaki tablo her fonksiyon için uygulanan kısıtlamayı göstermektedir:")

satirlar = [
    ("Fonksiyon", "Kısıtlama", "Gerekçe", "Durum"),
    ("submitTransaction", "onlyOwner", "Yalnızca kayıtlı sahipler teklif açabilir", "✅"),
    ("confirmTransaction", "onlyOwner + notConfirmed", "Duplikat onayı engeller", "✅"),
    ("executeTransaction", "onlyOwner + timeLockPassed", "Zaman kilidi ve yetki kontrolü", "✅"),
    ("revokeConfirmation", "onlyOwner + alreadyConfirmed", "Var olmayan onayı geri çekmeyi engeller", "✅"),
    ("pause", "onlyOwner", "Herhangi bir sahip durdurabilir (kilitlenme önlemi)", "✅"),
    ("unpause", "onlyOwner", "Herhangi bir sahip devam ettirebilir", "✅"),
    ("addOwner", "onlyWallet", "Sadece kontrat kendisi çağırabilir (multi-sig ile)", "✅"),
    ("removeOwner", "onlyWallet", "Sadece kontrat kendisi çağırabilir", "✅"),
    ("replaceOwner", "onlyWallet", "Sadece kontrat kendisi çağırabilir", "✅"),
    ("changeRequirement", "onlyWallet", "M değişikliği konsensüs gerektirir", "✅"),
    ("changeTimeLockDuration", "onlyWallet", "Süre değişikliği konsensüs gerektirir", "✅"),
]
tbl = doc.add_table(rows=len(satirlar), cols=4)
tbl.style = 'Table Grid'
for i, (a, b, c, d) in enumerate(satirlar):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
    tbl.rows[i].cells[2].text = c
    tbl.rows[i].cells[3].text = d
    if i == 0:
        for cell in tbl.rows[i].cells:
            set_cell_bg(cell, '185FA5')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            cell.paragraphs[0].runs[0].font.bold = True
    else:
        set_cell_bg(tbl.rows[i].cells[0], 'F0F7FF')
        tbl.rows[i].cells[3].paragraphs[0].runs[0].font.color.rgb = DARK_GREEN
        tbl.rows[i].cells[3].paragraphs[0].runs[0].font.bold = True
    for cell in tbl.rows[i].cells:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)

doc.add_paragraph()
doc.add_page_break()

add_heading(doc, "3.3 Saldırı Vektörü → Mitigasyon Tablosu", level=2)

add_body(doc,
    "Aşağıdaki tablo, çok imzalı cüzdan sistemlerine yönelik bilinen saldırı vektörlerini ve "
    "bu projede her birine karşı uygulanan savunma mekanizmalarını özetlemektedir:")

vektor = [
    ("Saldırı", "Vektör Açıklaması", "Mitigasyon", "Durum"),
    ("Reentrancy",
     "executeTransaction sırasında kötü niyetli bir kontrat geri çağrı yaparak işlemi tekrarlayabilir.",
     "OpenZeppelin ReentrancyGuard (nonReentrant modifier) + CEI pattern: executed=true nonce++ dış çağrıdan önce yazılır.",
     "✅ Korumalı"),
    ("Zincir İçi Replay",
     "Aynı txIndex iki kez yürütülmeye çalışılır.",
     "executed boolean flag: true olduktan sonra tekrar execute edilemez.",
     "✅ Korumalı"),
    ("Zincirler Arası Replay",
     "Sepolia üzerindeki bir TX'in imzası mainnet'e kopyalanır.",
     "DEPLOY_CHAIN_ID immutable değişkeni ve EIP-712 domain separator içindeki chainId.",
     "✅ Korumalı"),
    ("İmza Seti Replay",
     "Aynı off-chain imza seti farklı bir işlem için yeniden kullanılır.",
     "nonce: Her executeTransaction'da artar; DOMAIN_SEPARATOR içinde txIndex de kullanılır.",
     "✅ Korumalı"),
    ("Azınlık Saldırısı",
     "M'den az sahip konsensüs olmadan işlem yürütmeye çalışır.",
     "M ≥ N/2 hard-coded kuralı: constructor ve changeRequirement'da doğrulanır.",
     "✅ Korumalı"),
    ("Duplikat Onay",
     "Aynı sahip aynı işlemi birden fazla kez onaylar.",
     "confirmed[txIndex][owner] mapping + notConfirmed modifier.",
     "✅ Korumalı"),
    ("Time-lock Bypass",
     "M onaya ulaşır ulaşmaz anında execute edilmek istenir.",
     "timeLockPassed modifier: block.timestamp >= readyTime[txIndex] kontrolü.",
     "✅ Korumalı"),
    ("Yetkisiz Sahip Ekleme",
     "Bir sahip doğrudan addOwner() çağırarak yetkisiz adres ekler.",
     "onlyWallet modifier: Yalnızca kontratın kendisi (multi-sig sonrası) çağırabilir.",
     "✅ Korumalı"),
    ("Sıfır Adres Sahip",
     "address(0) sahip olarak eklenerek token yakılır veya erişilemez hale gelir.",
     "notNull modifier constructor ve addOwner'da kontrol eder.",
     "✅ Korumalı"),
    ("Duplikat Sahip",
     "Aynı adres iki kez eklenerek oy ağırlığı yapay olarak artırılır.",
     "ownerDoesNotExist modifier + isOwner mapping.",
     "✅ Korumalı"),
]
tbl = doc.add_table(rows=len(vektor), cols=4)
tbl.style = 'Table Grid'
col_widths = [Cm(3.2), Cm(5.5), Cm(6.5), Cm(1.8)]
for i, row in enumerate(tbl.rows):
    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]

for i, (a, b, c, d) in enumerate(vektor):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
    tbl.rows[i].cells[2].text = c
    tbl.rows[i].cells[3].text = d
    if i == 0:
        for cell in tbl.rows[i].cells:
            set_cell_bg(cell, '185FA5')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            cell.paragraphs[0].runs[0].font.bold = True
    else:
        set_cell_bg(tbl.rows[i].cells[0], 'FFF2CC')
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        tbl.rows[i].cells[3].paragraphs[0].runs[0].font.color.rgb = DARK_GREEN
        tbl.rows[i].cells[3].paragraphs[0].runs[0].font.bold = True
    for cell in tbl.rows[i].cells:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

add_heading(doc, "3.4 Güvenlik Mekanizmalarının Etkileşimi", level=2)

add_body(doc,
    "Güvenlik mekanizmaları birbirini tamamlayan bir savunma katmanı oluşturur (Defense in Depth) [11]. "
    "Bir mekanizmanın başarısız olması durumunda diğerleri devreye girer:")

for item in [
    "ReentrancyGuard → CEI → executed flag: Reentrancy saldırısına karşı üç bağımsız katman.",
    "chainId kontrolü → DOMAIN_SEPARATOR → nonce: Replay saldırısına karşı üç bağımsız katman.",
    "onlyOwner → notConfirmed → M eşiği: Yetkisiz onaya karşı üç bağımsız kontrol.",
    "emergencyPause: Tüm state-changing işlemleri dondurur; üstteki tüm mekanizmaları geçersiz kılar.",
]:
    add_bullet(doc, item)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════
# KAYNAKLAR
# ════════════════════════════════════════════════════════════
add_heading(doc, "Kaynaklar", level=1)

refs = [
    ('1', 'Nakamoto, S. (2008). Bitcoin: A Peer-to-Peer Electronic Cash System. bitcoin.org/bitcoin.pdf'),
    ('2', "Nathaniel Popper. (2014). In the Ruins of Bitcoin, a Puzzle: Mt. Gox's Missing Millions. The New York Times."),
    ('3', 'Jentzsch, C. (2016). Decentralized Autonomous Organization to Automate Governance. GitHub. github.com/blockchainsllc/DAO'),
    ('4', 'Gnosis Safe Team. (2023). Gnosis Safe: A Multisignature Wallet for Ethereum. docs.safe.global'),
    ('5', 'Martin, R. C. (2003). Agile Software Development: Principles, Patterns, and Practices. Prentice Hall.'),
    ('6', 'OpenZeppelin. (2024). ReentrancyGuard — OpenZeppelin Docs. docs.openzeppelin.com/contracts/5.x/api/utils#ReentrancyGuard'),
    ('7', 'Ethereum Foundation. (2018). EIP-712: Typed structured data hashing and signing. eips.ethereum.org/EIPS/eip-712'),
    ('8', 'Consensys. (2023). Smart Contract Best Practices: Checks-Effects-Interactions. consensys.github.io/smart-contract-best-practices/development-recommendations/general/'),
    ('9', 'Trail of Bits. (2024). Slither: Static Analyzer for Solidity. github.com/crytic/slither'),
    ('10', 'Ethereum Foundation. (2023). Block Timestamp Manipulation. swcregistry.io/docs/SWC-116'),
    ('11', 'Howard, M., LeBlanc, D. (2003). Writing Secure Code, 2nd ed. Microsoft Press.'),
    ('12', 'Solidity Documentation. (2024). Solidity v0.8.20. docs.soliditylang.org'),
    ('13', 'Ethereum Yellow Paper. (2024). ETHEREUM: A SECURE DECENTRALISED GENERALISED TRANSACTION LEDGER. ethereum.github.io/yellowpaper/paper.pdf'),
]
for num, text in refs:
    add_ref(doc, num, text)

# ── Kaydet ──────────────────────────────────────────────────
yol = r"c:\Users\merve\Desktop\8. dönem\blockchain\blockchain-prohe\Merve_Rapor_Bolum1_3.docx"
doc.save(yol)
print(f"Kaydedildi: {yol}")
