"""
MultiSigWallet — Kapsamlı Proje Raporu (Sıfırdan Anlatan)
Çalıştır: python scripts/generate_full_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Sayfa ayarları ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Renkler ───────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x18, 0x5F, 0xA5)
GREEN      = RGBColor(0x0F, 0x6E, 0x56)
ORANGE     = RGBColor(0x85, 0x4F, 0x0B)
PURPLE     = RGBColor(0x53, 0x4A, 0xB7)
DARK       = RGBColor(0x1A, 0x20, 0x2C)
MUTED      = RGBColor(0x71, 0x80, 0x96)
RED        = RGBColor(0x85, 0x1B, 0x1B)
DARK_GREEN = RGBColor(0x16, 0x65, 0x34)
LIGHT_BLUE = RGBColor(0x0E, 0x7E, 0xD4)

# ── Yardımcılar ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading(text, level=1, color=DARK):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
        run.font.size = Pt({1:17, 2:13, 3:11, 4:10}.get(level, 10))
    p.paragraph_format.space_before = Pt({1:18, 2:12, 3:8, 4:6}.get(level, 6))
    p.paragraph_format.space_after  = Pt(5)
    return p

def para(text="", bold=False, italic=False, size=10.5, color=DARK,
         space_after=5, indent=0, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Cm(indent)
    return p

def mixed_para(parts, size=10.5, space_after=5, indent=0):
    """parts = [(text, bold, color), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Cm(indent)
    for text, bold, color in parts:
        run = p.add_run(text)
        run.font.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
    return p

def code(text, indent=0.6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(indent)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EEF2F7")
    pPr.append(shd)
    return p

def bullet(parts_or_text, indent=0.5, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    if isinstance(parts_or_text, str):
        run = p.add_run(parts_or_text)
        run.font.size = Pt(size)
        run.font.color.rgb = DARK
    else:
        for text, bold, color in parts_or_text:
            run = p.add_run(text)
            run.font.bold = bold
            run.font.size = Pt(size)
            run.font.color.rgb = color
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    return p

def table(headers, rows, hdr_bg="185FA5", col_widths=None, hdr_size=9, row_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        set_cell_bg(c, hdr_bg)
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(hdr_size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, rdata in enumerate(rows):
        row = t.add_row()
        for ci, cell_text in enumerate(rdata):
            c = row.cells[ci]
            if ri % 2 == 1:
                set_cell_bg(c, "F5F8FC")
            r = c.paragraphs[0].add_run(str(cell_text))
            r.font.size = Pt(row_size)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return t

def info_box(title, text, bg_hex, border_color_rgb, text_color=DARK):
    p0 = doc.add_paragraph()
    r0 = p0.add_run(f"  {title}")
    r0.font.bold = True
    r0.font.size = Pt(9.5)
    r0.font.color.rgb = border_color_rgb
    pPr0 = p0._p.get_or_add_pPr()
    shd0 = OxmlElement("w:shd")
    shd0.set(qn("w:val"), "clear")
    shd0.set(qn("w:color"), "auto")
    shd0.set(qn("w:fill"), bg_hex)
    pPr0.append(shd0)
    p0.paragraph_format.space_after = Pt(0)

    p1 = doc.add_paragraph()
    r1 = p1.add_run(f"  {text}")
    r1.font.size = Pt(10)
    r1.font.color.rgb = text_color
    pPr1 = p1._p.get_or_add_pPr()
    shd1 = OxmlElement("w:shd")
    shd1.set(qn("w:val"), "clear")
    shd1.set(qn("w:color"), "auto")
    shd1.set(qn("w:fill"), bg_hex)
    pPr1.append(shd1)
    p1.paragraph_format.space_after = Pt(8)
    p1.paragraph_format.left_indent = Cm(0.3)

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "C7D9F0")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)


# ══════════════════════════════════════════════════════════════════════════════
#  KAPAK
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Multi-Signature Wallet")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Proje Raporu — Sıfırdan Sona Her Şey")
r2.font.size = Pt(14)
r2.font.color.rgb = DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("COM 4532 · Blockchain · Merve Çakır, Emine Binay, İrem Keskin · 2026")
r3.font.size = Pt(10)
r3.font.color.rgb = MUTED

doc.add_paragraph()

# Kısa içindekiler
para("Bu raporda neler var?", bold=True, size=11)
items = [
    "Bölüm 1 — Blockchain ve Ethereum nedir? (temel kavramlar)",
    "Bölüm 2 — Multi-Signature Wallet nedir, neden yapıyoruz?",
    "Bölüm 3 — Projenin mimarisi ve dosya yapısı",
    "Bölüm 4 — Solidity ve Smart Contract nedir?",
    "Bölüm 5 — Kontratlar: Ne yaptılar, nasıl çalışıyorlar?",
    "Bölüm 6 — Güvenlik mekanizmaları (neden lazımlar?)",
    "Bölüm 7 — Güvenlik audit: Slither ile ne bulduk?",
    "Bölüm 8 — Deploy: Blockchain'e yüklemek ne demek?",
    "Bölüm 9 — Emine için: Backend entegrasyonu",
    "Bölüm 10 — İrem için: Test ve UI",
]
for item in items:
    bullet(item)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 1 — BLOCKCHAIN VE ETHEREUM NEDİR?
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 1 — Blockchain ve Ethereum Nedir?", level=1, color=BLUE)
para(
    "Projeyi anlamak için önce birkaç temel kavramı bilmek gerekiyor. "
    "Blockchain biliyorsan bu bölümü atlayabilirsin.",
    size=10, color=MUTED, italic=True
)

heading("1.1 Blockchain Nedir?", level=2, color=DARK)
para(
    "Blockchain, yani 'blok zinciri', binlerce bilgisayarda aynı anda saklanan "
    "ve hiçbir merkezi otorite olmadan çalışan bir veritabanıdır. "
    "Normal bir bankada 'sen X'e Y TL gönderdin' bilgisi sadece bankanın "
    "sunucusunda tutulur. Blockchain'de ise bu bilgi binlerce farklı "
    "bilgisayarda aynı anda kayıtlıdır ve geriye dönük olarak değiştirilemez."
)
para(
    "Blockchain'in üç temel özelliği:"
)
bullet([("Merkeziyetsiz: ", True, DARK), ("Tek bir şirketin ya da kişinin kontrolünde değil.", False, DARK)])
bullet([("Şeffaf: ", True, DARK), ("Herkes her işlemi görebilir.", False, DARK)])
bullet([("Değiştirilemez: ", True, DARK), ("Bir kez yazılan kayıt silinemez ya da düzenlenemez.", False, DARK)])

heading("1.2 Ethereum Nedir?", level=2, color=DARK)
para(
    "Ethereum, Bitcoin gibi bir blockchain ağıdır ama üstüne bir şey daha ekler: "
    "Smart Contract (Akıllı Sözleşme). Ethereum'da sadece para göndermekle kalmaz, "
    "kod da çalıştırabilirsin. Bu sayede 'belirli koşullar gerçekleşince otomatik "
    "çalışan programlar' yazılabiliyor."
)
para(
    "ETH: Ethereum'un para birimi. İşlem ücretleri (gas) ETH ile ödenir."
)

heading("1.3 Gas Nedir?", level=2, color=DARK)
para(
    "Ethereum'da her işlem için bir ücret ödersin — buna gas denir. "
    "Ne kadar karmaşık bir işlem yaparsan (daha fazla hesaplama), "
    "o kadar fazla gas ödersin. Bu, ağın spam'e karşı korunma mekanizmasıdır."
)

heading("1.4 Testnet Nedir? (Sepolia)", level=2, color=DARK)
para(
    "Ethereum'un iki tür ağı var: Mainnet (gerçek para) ve Testnet (sahte para). "
    "Projemizde Sepolia adlı testnet'i kullanıyoruz. "
    "Buradaki ETH'nin gerçek değeri yok, faucet (musluk) adı verilen sitelerden "
    "ücretsiz alınabiliyor. Amacımız kodu gerçek parayı riske atmadan test etmek."
)

table(
    headers=["Kavram", "Açıklama", "Benzetme"],
    rows=[
        ["Blockchain", "Merkeziyetsiz, değiştirilemez defter", "Herkesin görebildiği, silemeceği bir not defteri"],
        ["Ethereum", "Üzerine program yazılabilen blockchain", "Programlanabilir not defteri"],
        ["ETH", "Ethereum'un para birimi", "Ağ içindeki para"],
        ["Gas", "İşlem ücreti", "Postane pulu — ne kadar şey göndersen o kadar ödersin"],
        ["Mainnet", "Gerçek Ethereum ağı", "Gerçek banka"],
        ["Testnet (Sepolia)", "Test amaçlı Ethereum ağı", "Eğitim kumbarası — sahte para ile pratik"],
        ["Smart Contract", "Blockchain'de çalışan otomatik program", "Otomatik işleyen sözleşme makinesi"],
    ],
    hdr_bg="185FA5",
    col_widths=[4, 6.5, 7.5],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 2 — MULTI-SIGNATURE WALLET NEDİR?
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 2 — Multi-Signature Wallet Nedir, Neden Yapıyoruz?", level=1, color=BLUE)

heading("2.1 Normal Kripto Cüzdanı", level=2, color=DARK)
para(
    "Normal bir kripto cüzdanında (MetaMask gibi) tek bir private key vardır. "
    "Bu key'e sahip olan kişi cüzdandaki her şeyi tamamen kontrol eder. "
    "Key çalınırsa ya da kaybolursa, cüzdandaki tüm varlıklar kaybedilir."
)
info_box(
    "Problem:",
    "Bir şirketin kasasında 1 milyon dolar ETH var ve bu kasanın tek bir anahtarı "
    "var. O anahtarı taşıyan kişi hastalanırsa? Anahtar çalınırsa? "
    "Ya da o kişi kötü niyetli çıkarsa?",
    "FEF2F2", RED, RED
)

heading("2.2 Multi-Signature (Çok İmzalı) Cüzdan", level=2, color=DARK)
para(
    "Multi-sig cüzdan, bir işlemin gerçekleşmesi için birden fazla kişinin onayını "
    "gerektirir. Tıpkı eski banka kasalarındaki gibi: kasayı açmak için hem banka "
    "müdürünün hem de müşterinin anahtarı gerekir."
)
para(
    "Bizim projemiz M-of-N modelini kullanıyor:"
)
bullet([("N:", True, DARK), (" Toplam sahip (owner) sayısı", False, DARK)])
bullet([("M:", True, DARK), (" Bir işlemin yürütülmesi için gereken minimum onay sayısı", False, DARK)])
bullet([("Kural:", True, DARK), (" M ≥ N/2 (çoğunluk onayı zorunlu)", False, DARK)])

para("Örnek senaryolar:", bold=True, size=10.5, space_after=3)
table(
    headers=["Yapılandırma", "Anlamı", "Kullanım Örneği"],
    rows=[
        ["2-of-3", "3 sahibin 2'si onaylamalı", "Küçük şirket kasası — bizim senaryo"],
        ["3-of-5", "5 sahibin 3'ü onaylamalı", "DAO (merkeziyetsiz org) hazinesi"],
        ["4-of-7", "7 sahibin 4'ü onaylamalı", "Büyük kurumsal cüzdan"],
    ],
    hdr_bg="185FA5",
    col_widths=[4, 6, 8.5],
)

heading("2.3 Neden Bu Projeyi Yapıyoruz?", level=2, color=DARK)
para(
    "COM 4532 Blockchain dersi kapsamında, gerçek dünyada kullanılan "
    "Multi-Signature Wallet mantığını sıfırdan implement ediyoruz. "
    "Gnosis Safe gibi endüstri standardı cüzdanların temel çalışma prensibini "
    "öğrenmek ve Solidity, Web3.py, testnet deploy gibi pratik becerileri "
    "kazanmak hedefleniyor."
)

heading("2.4 İşlem Akışı — Adım Adım", level=2, color=DARK)
para("Bir işlemin hayata geçmesi için şu adımlar yaşanır:", space_after=3)

steps = [
    ("1. Submit (Teklif)", "Sahiplerden biri 'şu adrese şu kadar ETH gönderelim' diye teklif sunar."),
    ("2. Confirm (Onay)", "Diğer sahipler bu teklifi görür ve tek tek onay verir."),
    ("3. Time-lock", "M onaya ulaşılınca 24 saatlik bir bekleme başlar (güvenlik)."),
    ("4. Execute (Yürüt)", "24 saat dolduktan sonra işlem zincirde gerçekleştirilir."),
]
table(
    headers=["Adım", "Ne Olur?"],
    rows=steps,
    hdr_bg="534AB7",
    col_widths=[4, 14.5],
)

info_box(
    "Gerçek Hayat Örneği:",
    "Merve, Emine ve İrem ortak bir kripto kasası işletiyor (3 sahip, M=2). "
    "Merve 'şu adrese 1 ETH gönderelim' teklifini sunar. "
    "Emine onaylar (1/2). İrem de onaylar (2/2 — yeter!). "
    "24 saat beklenip işlem yürütülür. "
    "Sadece Merve tek başına para gönderemez.",
    "E6F1FB", BLUE
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 3 — MİMARİ VE DOSYA YAPISI
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 3 — Projenin Mimarisi ve Dosya Yapısı", level=1, color=BLUE)

heading("3.1 Ekip İş Bölümü", level=2, color=DARK)
table(
    headers=["Kişi", "Sorumluluk", "Teknoloji"],
    rows=[
        ["Merve Çakır",  "Smart Contract yazımı, güvenlik audit, testnet deploy", "Solidity, Hardhat, Slither"],
        ["Emine Binay",  "Backend servisi, EIP-712 imza toplama, blockchain bağlantısı", "Python, Web3.py"],
        ["İrem Keskin",  "Kullanıcı arayüzü, birim testler, entegrasyon testleri", "Streamlit, Pytest"],
    ],
    hdr_bg="534AB7",
    col_widths=[4, 7.5, 7],
)

heading("3.2 Genel Mimari", level=2, color=DARK)
para("Sistem üç katmandan oluşuyor:", space_after=3)
table(
    headers=["Katman", "Ne Yapar?", "Kim Yazdı?"],
    rows=[
        ["Smart Contract (Blockchain)", "İşlemleri, onayları, sahipleri blockchain'de güvenli tutar", "Merve"],
        ["Backend (Python/Web3.py)",    "Kullanıcının yaptıklarını contract'a iletir, olayları dinler", "Emine"],
        ["Frontend (Streamlit)",        "Kullanıcının gördüğü ekran — butonlar, listeler, bildirimler", "İrem"],
    ],
    hdr_bg="185FA5",
    col_widths=[5, 8.5, 5],
)

heading("3.3 Contract Mimarisi", level=2, color=DARK)
para(
    "Projede 3 Solidity dosyası var ama deploy edilen (blockchain'e yüklenen) "
    "sadece bir tanesi: MultiSigWallet. Diğer ikisi 'abstract contract' — "
    "yani kendi başlarına çalışamayan, MultiSigWallet tarafından içine dahil edilen parçalar."
)
code("MultiSigWallet.sol          ← TEK deploy edilen contract")
code("  ├── is ReentrancyGuard    ← OpenZeppelin kütüphanesi (güvenlik)")
code("  ├── is OwnerManager.sol   ← abstract: sahip yönetimi")
code("  └── is TransactionManager.sol  ← abstract: işlem yaşam döngüsü")

para(
    "'Abstract contract' ne demek? Tıpkı bir şablona benzetebilirsin. "
    "OwnerManager.sol, 'sahip ekle/çıkar' mantığını tanımlar ama kendi başına "
    "çalışamaz. MultiSigWallet bu şablonu alır ve kendi içinde kullanır.",
    size=10, indent=0.3, space_after=8
)

heading("3.4 Dosya ve Klasör Yapısı", level=2, color=DARK)
table(
    headers=["Dosya / Klasör", "Ne İşe Yarar?", "Kim Kullanır?"],
    rows=[
        ["contracts/MultiSigWallet.sol", "Ana kontrat — tüm mantık burada birleşiyor", "Merve yazar"],
        ["contracts/OwnerManager.sol", "Sahip ekleme/çıkarma/değiştirme mantığı", "Merve yazar"],
        ["contracts/TransactionManager.sol", "İşlem struct'ı, onay takibi, time-lock state", "Merve yazar"],
        ["abi/MultiSigWallet.abi.json", "Kontratın 'dışarıya açık kapıları' listesi", "Emine kullanır"],
        ["scripts/deploy.js", "Kontratı Sepolia'ya yükleyen script", "Merve çalıştırır"],
        ["scripts/verify.js", "Etherscan'de kaynak kodu doğrulama", "Merve çalıştırır"],
        ["security_audit.md", "Slither güvenlik tarama raporu", "Raporda kullanılır"],
        [".env.example", "Gerekli ayarların şablonu (.env'ye kopyalanır)", "Herkes doldurur"],
    ],
    hdr_bg="185FA5",
    col_widths=[5.5, 7, 4.5],
)

info_box(
    "ABI Nedir?",
    "ABI (Application Binary Interface), bir kontratın dışarıdan çağrılabilen "
    "fonksiyonlarının listesidir. Emine Python'dan 'submitTransaction' çağırmak "
    "istediğinde, Web3.py bu ABI dosyasına bakarak nasıl çağıracağını anlar. "
    "ABI = kontratın 'kullanım kılavuzu'.",
    "E6F1FB", BLUE
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 4 — SOLİDİTY VE SMART CONTRACT NEDİR?
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 4 — Solidity ve Smart Contract Nedir?", level=1, color=BLUE)

heading("4.1 Smart Contract (Akıllı Sözleşme)", level=2, color=DARK)
para(
    "Smart contract, Ethereum blockchain üzerinde çalışan bir programdır. "
    "Ama normal bir programdan şu farkla: bir kez deploy edildiğinde "
    "kimse onu durduramaz, değiştiremez ve manipüle edemez. "
    "Kod ne diyorsa o olur."
)
para(
    "Örnek: 'Eğer 2 sahip onaylarsa, parayı gönder.' Bu kural kontrata "
    "yazıldıktan sonra ne Merve ne Emine ne de başka biri bu kuralı "
    "devre dışı bırakamaz. Kod otomatik çalışır."
)

heading("4.2 Solidity Nedir?", level=2, color=DARK)
para(
    "Solidity, Ethereum üzerinde smart contract yazmak için kullanılan "
    "programlama dilidir. JavaScript'e benzer bir sözdizimine sahiptir. "
    "Solidity kodu derlenir (compile) ve Ethereum Virtual Machine (EVM) "
    "üzerinde çalışır."
)

heading("4.3 Temel Solidity Kavramları", level=2, color=DARK)
table(
    headers=["Kavram", "Ne Demek?", "Örnekte Ne Gördük?"],
    rows=[
        ["contract", "Smart contract'ın kendisi — bir sınıf gibi", "contract MultiSigWallet { }"],
        ["function", "Çağrılabilir işlem — kullanıcı veya başka contract çağırır", "function submitTransaction(...)"],
        ["mapping", "Key-value harita (sözlük) — çok hızlı sorgulama", "mapping(address => bool) isOwner"],
        ["event", "Blockchain'e kayıt düşülen bildirim — loglar", "event Deposit(address, uint256)"],
        ["modifier", "Fonksiyon çağrılmadan önce çalışan kontrol bloğu", "modifier onlyOwner()"],
        ["require", "Koşul sağlanmazsa işlemi durdur ve hata mesajı ver", "require(M >= N/2, '...')"],
        ["public", "Her yerden çağrılabilir", "function getBalance() public"],
        ["view", "Sadece okur, blockchain'i değiştirmez", "function isConfirmed() public view"],
        ["immutable", "Sadece constructor'da set edilebilir, sonra değiştirilemez", "uint256 immutable CHAIN_ID"],
        ["payable", "ETH alabilir", "receive() external payable"],
    ],
    hdr_bg="534AB7",
    col_widths=[3.5, 6, 8.5],
)

heading("4.4 OpenZeppelin Nedir?", level=2, color=DARK)
para(
    "OpenZeppelin, güvenliği denetlenmiş hazır Solidity kütüphaneleri sunan "
    "bir projedir. Tıpkı Python'da numpy veya pandas gibi — tekerleği yeniden "
    "icat etmene gerek kalmaz, test edilmiş güvenli kodu kullanırsın. "
    "Biz projemizde OpenZeppelin'in ReentrancyGuard'ını kullandık."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 5 — KONTRATLAR: NE YAPIYORLAR?
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 5 — Kontratlar: Ne Yaptılar, Nasıl Çalışıyorlar?", level=1, color=BLUE)

# 5.1 OwnerManager
heading("5.1 OwnerManager.sol — Sahip Yönetimi", level=2, color=DARK)
para(
    "Bu kontrat, 'kim bu cüzdanın sahibidir?' sorusunu yönetir. "
    "Sahip eklemek, çıkarmak, değiştirmek ve onay eşiğini güncellemek "
    "bu kontratın işidir."
)
para("Içinde ne var?", bold=True, size=10.5, space_after=3)
table(
    headers=["Ne?", "Türü", "Ne İşe Yarar?"],
    rows=[
        ["isOwner", "mapping(address → bool)", "Bir adresin sahip olup olmadığını tutar — O(1) sorgulama"],
        ["owners", "address[] (dizi)", "Tüm sahiplerin sıralı listesi"],
        ["numConfirmationsRequired", "uint256 (sayı)", "Bir işlemin yürütülmesi için gereken minimum onay (M)"],
        ["addOwner(address)", "Fonksiyon", "Yeni sahip ekler — M ≥ N/2 kuralını korur"],
        ["removeOwner(address)", "Fonksiyon", "Sahibi çıkarır — gerekirse M'i otomatik düşürür"],
        ["replaceOwner(eski, yeni)", "Fonksiyon", "Bir sahibi başka biriyle değiştirir"],
        ["changeRequirement(M)", "Fonksiyon", "Onay eşiğini günceller"],
        ["getOwners()", "Fonksiyon", "Tüm sahiplerin listesini döndürür"],
    ],
    hdr_bg="185FA5",
    col_widths=[4.5, 4.5, 9.5],
)
info_box(
    "onlyWallet ne demek?",
    "addOwner, removeOwner gibi fonksiyonlar 'onlyWallet' modifier'ına sahip. "
    "Bu demek ki sadece kontratın kendisi bu fonksiyonları çağırabilir. "
    "Yani Merve bile doğrudan 'addOwner çağrayayım' diyemez. "
    "Önce bir teklif oluşturulur, M onay alınır, sonra kontrat kendi kendini çağırır. "
    "Bu sayede tek kişi sistemi ele geçiremiyor.",
    "E6F1FB", BLUE
)

info_box(
    "M ≥ N/2 Kuralı Neden Var?",
    "Diyelim ki 4 sahibin var ve M=1 yaptın. O zaman herhangi bir sahip "
    "tek başına kasa boşaltabilir — multi-sig'in anlamı kalmaz. "
    "Bu yüzden onay sayısı en az sahip sayısının yarısı kadar olmalı. "
    "3 sahip varsa M en az 2 olmalı. Bu kural her sahip değişiminde yeniden kontrol edilir.",
    "FFF7ED", ORANGE, ORANGE
)

divider()

# 5.2 TransactionManager
heading("5.2 TransactionManager.sol — İşlem Yaşam Döngüsü", level=2, color=DARK)
para(
    "Bu kontrat, bir işlemin tüm verilerini ve durumunu tutar. "
    "Hangi işlem bekliyor, kim onayladı, time-lock ne zaman bitiyor — "
    "hepsini bu kontrat saklıyor."
)
para("Transaction Struct — Her İşlemin Kaydı:", bold=True, size=10.5, space_after=3)
code("struct Transaction {")
code("    address to;               // Parayı alacak adres")
code("    uint256 value;            // Gönderilecek ETH miktarı (en küçük birim: wei)")
code("    bytes data;               // Fonksiyon çağrısı verisi (boş da olabilir)")
code("    bool executed;            // Bu işlem zaten yapıldı mı?")
code("    uint256 numConfirmations; // Kaç sahip onayladı?")
code("}")
doc.add_paragraph()

info_box(
    "Wei Nedir?",
    "1 ETH = 1,000,000,000,000,000,000 wei (10^18). "
    "Smart contract'larda para her zaman wei cinsinden tutulur — "
    "küsuratlı hesaplama hatalarını önlemek için.",
    "E6F1FB", BLUE
)

para("State değişkenleri (blockchain'de saklanan veriler):", bold=True, size=10.5, space_after=3)
table(
    headers=["Değişken", "Türü", "Ne Saklar?"],
    rows=[
        ["transactions", "Transaction[]", "Tüm işlem kayıtları — txIndex ile erişilir"],
        ["confirmed", "mapping(txIndex → mapping(adres → bool))", "Kim hangi işlemi onayladı?"],
        ["readyTime", "mapping(txIndex → uint256)", "İşlem ne zaman yürütülebilir? (time-lock)"],
        ["timeLockDuration", "uint256", "Bekleme süresi — varsayılan 86400 saniye (24 saat)"],
    ],
    hdr_bg="185FA5",
    col_widths=[4, 6, 8.5],
)

divider()

# 5.3 MultiSigWallet — Ana kontrat
heading("5.3 MultiSigWallet.sol — Ana Kontrat", level=2, color=DARK)
para(
    "Hem OwnerManager hem de TransactionManager'ı miras alan, "
    "deploy edilen tek kontrat. Tüm core fonksiyonlar burada."
)

heading("Constructor — Cüzdanı Başlatma", level=3, color=DARK)
para(
    "Bir kontrat deploy edildiğinde ilk çalışan fonksiyondur. "
    "Bir kez çalışır ve bir daha çalışmaz. Sahipleri ve M değerini belirler."
)
code("constructor(address[] owners, uint256 M)")
code("  ├── owners boş olamaz")
code("  ├── M ≥ N/2 zorunlu (ihlal edilirse deploy başarısız)")
code("  ├── Aynı adres iki kez verilemez")
code("  ├── address(0) (sıfır adres) verilemez")
code("  └── EIP-712 domain separator hesaplanır (replay koruması için)")

doc.add_paragraph()
heading("submitTransaction — İşlem Teklifi Oluşturma", level=3, color=DARK)
para(
    "Sahiplerden biri 'şunu yapalım' diye teklif sunar. "
    "Bu fonksiyon çağrıldığında işlem 'pending' (bekleyen) listesine eklenir "
    "ve bir txIndex (işlem numarası) döner."
)
code("submitTransaction(address hedef, uint256 miktar, bytes veri)")
code("  → txIndex döner (örn: 0, 1, 2...)")
para(
    "Sadece sahipler çağırabilir. Hedef adres sıfır olamaz. "
    "Cüzdan dondurulmuşsa (paused) çalışmaz.",
    size=10, indent=0.3, color=MUTED
)

doc.add_paragraph()
heading("confirmTransaction — Onay Verme", level=3, color=DARK)
para(
    "Bir sahip bekleyen işlemi onaylar. "
    "Aynı sahip aynı işlemi iki kez onaylayamaz. "
    "Onay sayısı M'e ulaşınca time-lock başlar (24 saat sayaç)."
)
code("confirmTransaction(txIndex)")
para("Zaten yürütülmüş işlem onaylanamaz. Cüzdan dondurulmuşsa çalışmaz.", size=10, indent=0.3, color=MUTED)

doc.add_paragraph()
heading("executeTransaction — İşlemi Yürütme", level=3, color=DARK)
para(
    "M onay alınmış ve 24 saat dolmuşsa, herhangi bir sahip bu fonksiyonu "
    "çağırarak işlemi zincirde gerçekleştirir. Para aktarılır."
)
code("executeTransaction(txIndex)")
para("M onay YOK → revert. 24 saat dolmamış → revert. Zaten yürütülmüş → revert.", size=10, indent=0.3, color=MUTED)

doc.add_paragraph()
heading("revokeConfirmation — Onayı Geri Çekme", level=3, color=DARK)
para(
    "Bir sahip daha önce verdiği onaydan vazgeçebilir. "
    "Sadece henüz yürütülmemiş işlemler için geçerlidir."
)
code("revokeConfirmation(txIndex)")

doc.add_paragraph()
heading("Sorgu Fonksiyonları (View)", level=3, color=DARK)
para("Bunlar blockchini değiştirmez, sadece okur — gas ücreti çok düşük ya da sıfır:", space_after=3)
table(
    headers=["Fonksiyon", "Ne Döndürür?", "Kullanım Amacı"],
    rows=[
        ["getTransaction(txIndex)", "to, value, data, executed, numConfirmations", "İşlem detaylarını göster"],
        ["getTransactionCount()", "Toplam işlem sayısı", "Tüm işlemleri listelemek için döngü"],
        ["getConfirmationStatus(txIndex, owner)", "bool (onayladı mı?)", "Kimin onayladığını göster"],
        ["getOwners()", "address[] (tüm sahipler)", "Dashboard'da sahip listesi"],
        ["isConfirmed(txIndex)", "bool (M onaya ulaştı mı?)", "Execute edilebilir mi?"],
        ["getBalance()", "uint256 (wei)", "Cüzdan ETH bakiyesi"],
        ["getChainId()", "uint256", "Hangi blockchain'deyiz?"],
        ["getTransactionHash(txIndex)", "bytes32", "EIP-712 imza için hash"],
    ],
    hdr_bg="185FA5",
    col_widths=[5, 5.5, 7],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 6 — GÜVENLİK MEKANİZMALARI
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 6 — Güvenlik Mekanizmaları", level=1, color=BLUE)
para(
    "Blockchain'de bir hata veya güvenlik açığı çok ciddi sonuçlar doğurabilir: "
    "kontrat değiştirilemez, kayıplar geri alınamaz. Bu yüzden her olası "
    "saldırı senaryosuna karşı önlem alınması şart."
)

# 6.1 Reentrancy
heading("6.1 Reentrancy Saldırısı ve ReentrancyGuard", level=2, color=DARK)
para(
    "Reentrancy, blockchain'in en ünlü saldırı türü. "
    "2016 yılında bu saldırıyla 60 milyon dolar çalındı (The DAO hack)."
)
para("Saldırı nasıl çalışır?", bold=True, size=10.5, space_after=3)
bullet("Kötü amaçlı bir kontrat, bizim kontratımızdan para çekmeye çalışır.")
bullet("Para gönderilirken kötü kontrat 'ben parayı aldım, ama bir daha isteyeyim' der ve aynı fonksiyonu tekrar çağırır.")
bullet("Bakiye güncellenmeden önce ikinci çekme gerçekleşir — böylece para birden fazla kez çekilebilir.")
doc.add_paragraph()
para("Nasıl koruduk?", bold=True, size=10.5, space_after=3)
bullet([("ReentrancyGuard:", True, DARK), (" OpenZeppelin'den alınan bir 'kilit' mekanizması. Bir fonksiyon çalışırken aynı fonksiyon tekrar çağrılamaz.", False, DARK)])
bullet([("CEI Pattern:", True, DARK), (" Check-Effects-Interactions — önce durumu güncelle, sonra dış çağrı yap.", False, DARK)])
code("txn.executed = true;   // Önce 'yürütüldü' diye işaretle (CEI)")
code("nonce++;               // Sonra nonce'u artır (CEI)")
code("txn.to.call{...}();   // En son dış çağrıyı yap")
para("Böylece saldırgan geri çağırırsa 'zaten yürütüldü' hatası alır.", size=10, indent=0.3, color=MUTED)

divider()

# 6.2 Time-lock
heading("6.2 Time-lock (24 Saat Bekleme)", level=2, color=DARK)
para(
    "M onay alındıktan sonra işlem hemen yürütülmez — 24 saat bekler. "
    "Bu süre içinde eğer bir şeyler yanlış gidiyorsa (örn. bir sahip hacklenmiş "
    "ve onayı zorla alınmışsa), diğer sahipler fark edip onaylarını geri çekebilir."
)
para("Teknik detay:", bold=True, size=10.5, space_after=3)
bullet("M. onay verildiğinde → readyTime[txIndex] = şimdiki zaman + 24 saat olarak set edilir")
bullet("TimeLockTriggered event yayınlanır — Emine bu event'i dinleyip UI'a 'X saat sonra yürütülebilir' yazabilir")
bullet("executeTransaction → block.timestamp >= readyTime[txIndex] kontrolü yapar")
bullet("24 saat dolmamışsa 'time-lock suresi dolmadi' hatasıyla işlem durur")
doc.add_paragraph()
info_box(
    "Önemli:",
    "Timer bir kez başlayınca sıfırlanmaz. Bir sahip onayını geri çekip "
    "yeniden verse de 24 saatlik sayaç ilk M onay anından itibaren işler.",
    "FFF7ED", ORANGE
)

divider()

# 6.3 emergencyPause
heading("6.3 emergencyPause (Acil Durdurma)", level=2, color=DARK)
para(
    "Beklenmedik bir güvenlik açığı keşfedilirse, herhangi bir sahip "
    "cüzdanı tek başına dondurabilir. Dondurulunca submit, confirm, execute, "
    "revoke — hiçbiri çalışmaz."
)
para("Neden tek sahip dondurabilir?", bold=True, size=10.5, space_after=3)
para(
    "Acil durumlarda zaman kritiktir. Saldırı devam ederken 3 sahibin "
    "toplanıp multisig kararı alması çok zaman alır. Hız için tek sahip yetki verildi.",
    indent=0.3
)
para("Neden açmak da tek sahiple mümkün?", bold=True, size=10.5, space_after=3)
para(
    "Cüzdan dondurulunca executeTransaction çalışmıyor. Eğer açmak için "
    "multisig konsensüsü gerekseydi, açmak için execute gerekir ama execute "
    "çalışmıyor — sonsuza kadar kilitli kalırdı (deadlock). Bu yüzden açmak da "
    "tek sahip kararıyla yapılabiliyor.",
    indent=0.3
)

divider()

# 6.4 Replay
heading("6.4 Replay Koruması (EIP-712 + Nonce + Chain ID)", level=2, color=DARK)
para("İki farklı replay (tekrarlama) saldırısına karşı önlem alındı:", space_after=3)

para("Saldırı 1 — Zincirler Arası Replay:", bold=True, size=10.5, space_after=3)
bullet("Sepolia testnet'inde yapılan bir işlem, aynı adresler Mainnet'te de varsa tekrarlanabilir.")
bullet([("Çözüm:", True, DARK), (" DOMAIN_SEPARATOR — chainId (zincir kimliği) ve kontrat adresini içerir. Sepolia'daki imza Mainnet'te geçersizdir.", False, DARK)])

doc.add_paragraph()
para("Saldırı 2 — İmza Seti Replay:", bold=True, size=10.5, space_after=3)
bullet("Emine off-chain imzaları toplar ve işlemi yürütür. Saldırgan bu imzaları kaydeder ve bir sonraki fırsatta aynı imzaları tekrar kullanmaya çalışır.")
bullet([("Çözüm:", True, DARK), (" nonce — her executeTransaction'da 1 artar. Eski imzalar yeni nonce ile uyuşmaz → geçersiz.", False, DARK)])

doc.add_paragraph()
para("Kontrata eklenen özellikler:", bold=True, size=10.5, space_after=3)
table(
    headers=["Özellik", "Ne Yapar?"],
    rows=[
        ["DOMAIN_SEPARATOR", "chainId + kontrat adresi hash'i — her zincir/kontrat için benzersiz"],
        ["nonce", "Her başarılı execute'da 1 artar — eski imzaları geçersiz kılar"],
        ["DEPLOY_CHAIN_ID", "Deploy anındaki chainId saklanır — yanlış zincirde execute durur"],
        ["getTransactionHash()", "EIP-712 uyumlu imzalanacak hash — Emine imza toplarken kullanır"],
    ],
    hdr_bg="185FA5",
    col_widths=[5, 13.5],
)

divider()

# 6.5 Genel güvenlik tablosu
heading("6.5 Tüm Güvenlik Mekanizmaları — Özet", level=2, color=DARK)
table(
    headers=["Mekanizma", "Neye Karşı Korur?", "Durum"],
    rows=[
        ["M ≥ N/2 hard-code", "Azınlık saldırısı — tek kişi karar alamaz", "✅"],
        ["ReentrancyGuard", "Reentrancy — aynı fonksiyonu tekrar çağırma", "✅"],
        ["CEI Pattern", "Reentrancy — durum güncellenmeden dış çağrı", "✅"],
        ["executed flag", "Aynı işlemi iki kez yürütme", "✅"],
        ["confirmed mapping", "Aynı sahip iki kez onay verme", "✅"],
        ["Time-lock (24h)", "Acele alınan kötü kararlar", "✅"],
        ["emergencyPause", "Fark edilen güvenlik açığında hızlı durdurma", "✅"],
        ["DOMAIN_SEPARATOR", "Zincirler arası replay saldırısı", "✅"],
        ["nonce", "İmza seti replay saldırısı", "✅"],
        ["DEPLOY_CHAIN_ID", "Chain fork/split sonrası yanlış zincirde işlem", "✅"],
        ["notNull modifier", "Sıfır adrese para gönderme hatası", "✅"],
    ],
    hdr_bg="0F6E56",
    col_widths=[5.5, 8.5, 2.5],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 7 — GÜVENLİK AUDIT: SLİTHER
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 7 — Güvenlik Audit: Slither ile Ne Bulduk?", level=1, color=BLUE)

heading("7.1 Slither Nedir?", level=2, color=DARK)
para(
    "Slither, Solidity kontratlarını otomatik olarak analiz eden ve "
    "potansiyel güvenlik açıklarını tespit eden bir araçtır. "
    "Trail of Bits şirketi tarafından geliştirilmiştir ve endüstri standardıdır. "
    "Kontratın kaynak kodunu ve derlenen bytecode'u analiz eder, "
    "100'den fazla farklı güvenlik deseni kontrol eder."
)

heading("7.2 Audit Süreci", level=2, color=DARK)
para("Slither kuruldu ve tüm kontratlar üzerinde çalıştırıldı:", space_after=3)
code("pip install slither-analyzer")
code("slither contracts/MultiSigWallet.sol")
doc.add_paragraph()
para("İki kez çalıştırıldı: ilk analiz ve düzeltmeler sonrası kontrol.", space_after=3)

heading("7.3 Sonuçlar", level=2, color=DARK)
table(
    headers=["Önem Seviyesi", "Açıklama", "İlk", "Sonrası"],
    rows=[
        ["Critical", "Anında istismar edilebilir kritik açık", "0", "0"],
        ["High",     "Ciddi güvenlik riski",                    "0", "0"],
        ["Medium",   "Orta düzey risk — düzeltilmeli",          "2", "0 ✅"],
        ["Low/Info", "Düşük risk veya bilgi amaçlı",            "5", "5*"],
    ],
    hdr_bg="185FA5",
    col_widths=[3.5, 8, 2.5, 2.5],
)
para("* 5 bulgunun tamamı false positive (OZ kütüphanesi) veya kasıtlı tasarım kararı.", size=9, color=MUTED, italic=True)

heading("7.4 Bulunan ve Düzeltilen Sorunlar", level=2, color=DARK)
para("[DÜZELTİLDİ] Strict Equality — Medium Seviye", bold=True, size=10.5)
para(
    "readyTime[txIndex] == 0 kontrolü riskli bulundu. "
    "Eğer çok uç bir senaryoda timestamp manipülasyonu olursa "
    "bu kontrol yanlış sonuç verebilirdi. "
    "Düzeltme: ayrı bir boolean flag (_timeLockInitialized) eklendi.",
    indent=0.3
)
doc.add_paragraph()
para("[DÜZELTİLDİ] Array Length Cache — Low (Gas Optimizasyonu)", bold=True, size=10.5)
para(
    "owners.length her döngü adımında storage'dan okunuyordu. "
    "Storage okuma Ethereum'da pahalıdır (gas tüketir). "
    "Düzeltme: döngüden önce uint256 len = owners.length ile cache'lendi. "
    "Küçük ama gerçek bir gas tasarrufu.",
    indent=0.3
)

heading("7.5 Kalan Bulgular — Neden Düzeltilmedi?", level=2, color=DARK)
table(
    headers=["Bulgu", "Neden Düzeltilmedi?"],
    rows=[
        ["Assembly kodu",
         "OpenZeppelin kütüphanesinin kendi kodu — biz yazmadık, değiştiremeyiz"],
        ["Pragma versiyonu",
         "OZ ^0.8.20 kullanıyor, bizim kontratlarımız 0.8.20 sabitlenmiş — conflict kaçınılmaz"],
        ["solc 0.8.20 bug'ları",
         "Bahsedilen 3 bug (VerbatimInvalidDeduplication vb.) projemizi etkilemiyor — kullanmadığımız pattern'lar"],
        ["Low-level call",
         "Multisig için zorunlu — hedef kontrat compile-time'da belli değil, yüksek seviye çağrı yapılamaz"],
        ["İsimlendirme",
         "_param prefix'i Solidity standardı. DOMAIN_SEPARATOR ALL_CAPS immutable için önerilen biçim"],
    ],
    hdr_bg="185FA5",
    col_widths=[4.5, 14],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 8 — DEPLOY: BLOCKCHAIN'E YÜKLEMEK NE DEMEK?
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 8 — Deploy: Blockchain'e Yüklemek Ne Demek?", level=1, color=BLUE)

heading("8.1 Deploy Nedir?", level=2, color=DARK)
para(
    "Solidity kodu yazıp derlediğimizde elimizde iki şey oluyor: "
    "ABI (fonksiyon listesi) ve bytecode (makinenin anlayacağı düşük seviye kod). "
    "Deploy, bu bytecode'u Ethereum blockchain'e göndermek ve "
    "kontrata kalıcı bir adres aldırmak demektir."
)
para(
    "Kontrat deploy edildikten sonra o adreste sonsuza kadar yaşar. "
    "Hiç kimse onu silemez ya da değiştiremez. "
    "Emine ve İrem bu adres üzerinden kontratı kullanır."
)
table(
    headers=["Aşama", "Ne Olur?"],
    rows=[
        ["Yazım", "Solidity kodu yazılır (bitti ✅)"],
        ["Derleme (Compile)", "Hardhat solidity kodunu bytecode'a çevirir (bitti ✅)"],
        ["Deploy", "Bytecode Sepolia'ya gönderilir, kontrat adresi alınır (yarın 26 Mayıs)"],
        ["Verify", "Kaynak kod Etherscan'e yüklenir — herkes 'bu gerçekten MultiSigWallet' diye görebilir"],
        ["Kullanım", "Emine backend'i bu adrese bağlar, İrem Streamlit üzerinden demo yapar"],
    ],
    hdr_bg="185FA5",
    col_widths=[4.5, 14],
)

heading("8.2 Hardhat Nedir?", level=2, color=DARK)
para(
    "Hardhat, Solidity geliştirme için kullanılan bir araç setidir. "
    "Kontratları derlemek (compile), local olarak test etmek ve "
    "gerçek ağlara deploy etmek için kullanılıyor. "
    "Bizim projemizde Hardhat ile tüm kontratlar derlendi."
)

heading("8.3 Deploy Nasıl Yapılacak? (26 Mayıs)", level=2, color=DARK)
para("Yarın Merve şu adımları yapacak:", space_after=3)
table(
    headers=["Adım", "Ne Yapılır?", "Nereden?"],
    rows=[
        ["1", "Sepolia ETH al (test parası)", "sepoliafaucet.com"],
        ["2", "Sepolia RPC URL al (ağa bağlanmak için)", "infura.io veya alchemy.com — ücretsiz"],
        ["3", "Etherscan API key al (verify için)", "etherscan.io → API Keys"],
        ["4", "Emine ve İrem'den adres al", "Yarınki toplantıda"],
        ["5", ".env dosyasını doldur", "Yukarıdaki bilgilerle"],
        ["6", "node scripts/deploy.js çalıştır", "Terminal"],
        ["7", "node scripts/verify.js çalıştır", "Terminal"],
        ["8", "Contract adresini Emine ve İrem'e ver", "Mesajla"],
    ],
    hdr_bg="185FA5",
    col_widths=[1.5, 8, 9],
)

heading("8.4 Etherscan Verify Ne Demek?", level=2, color=DARK)
para(
    "Bir kontrat deploy edildiğinde adresine gidersen sadece bytecode görürsün "
    "(0x6080604052... gibi sayılardan oluşan anlamsız bir şey). "
    "Verify işlemi, kaynak kodunu (okunabilir Solidity kodu) Etherscan'e yükler "
    "ve 'bu bytecode bu koddan geldi' diye onaylar. "
    "Böylece herkes kontratın ne yaptığını şeffaf biçimde görebilir."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 9 — EMİNE İÇİN: BACKEND ENTEGRASYONU
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 9 — Emine İçin: Backend Entegrasyonu", level=1, color=GREEN)
para(
    "Bu bölüm Emine'nin Web3.py backend'ini kontratla bağlaması için "
    "ihtiyaç duyduğu her şeyi içeriyor.",
    size=10, color=MUTED, italic=True
)

heading("9.1 Kullanacağın Tek ABI Dosyası", level=2, color=GREEN)
para(
    "OwnerManager ve TransactionManager'ın fonksiyonları MultiSigWallet ABI'sinin "
    "içinde zaten var. Sadece şunu kullan:"
)
code("abi/MultiSigWallet.abi.json")
code("")
code("# Python'da:")
code("import json")
code("with open('abi/MultiSigWallet.abi.json') as f:")
code("    abi = json.load(f)")
code("contract = w3.eth.contract(address=CONTRACT_ADDRESS, abi=abi)")

heading("9.2 İşlem Akışı — Python Kodunda", level=2, color=GREEN)
para("Tam 2-of-3 akışı:", bold=True, space_after=3)
code("# Adım 1: Merve işlem teklif eder")
code("tx_index = contract.functions.submitTransaction(")
code("    hedef_adres,   # to")
code("    wei_miktari,   # value (1 ETH = 10**18 wei)")
code("    b''            # data (boş = sadece ETH transfer)")
code(").transact({'from': merve_adresi})")
code("")
code("# Adım 2: Emine onaylar")
code("contract.functions.confirmTransaction(tx_index).transact({'from': emine_adresi})")
code("")
code("# Adım 3: İrem onaylar (M=2 tamamlandı → time-lock başladı)")
code("contract.functions.confirmTransaction(tx_index).transact({'from': irem_adresi})")
code("")
code("# Adım 4: 24 saat sonra herhangi biri execute eder")
code("contract.functions.executeTransaction(tx_index).transact({'from': emine_adresi})")

heading("9.3 Yeni Özelliklerin Entegrasyonu", level=2, color=GREEN)
table(
    headers=["Özellik", "Nasıl Kullanırsın?"],
    rows=[
        ["paused()",
         "İşlem göndermeden önce kontrol et — True ise UI'da 'Cüzdan dondurulmuş' uyarısı göster"],
        ["readyTime(txIndex)",
         "Unix timestamp döner. time.time() ile karşılaştır → kalan süreyi hesapla"],
        ["TimeLockTriggered event",
         "M onaya ulaşınca emit edilir. args['unlockTime'] → UI'da geri sayım göster"],
        ["nonce()",
         "EIP-712 off-chain imzalarda güncel nonce'u imza struct'ına dahil et"],
        ["getTransactionHash(txIndex)",
         "İmzalanacak hash'i kontraktan al — sahiplere imzalat"],
        ["DOMAIN_SEPARATOR()",
         "Off-chain imza doğrulamasında kullan"],
        ["EmergencyPaused event",
         "Dinle → UI'da anlık uyarı göster"],
        ["EmergencyUnpaused event",
         "Dinle → UI'da 'cüzdan aktif' güncelle"],
    ],
    hdr_bg="0F6E56",
    col_widths=[5, 13.5],
)

heading("9.4 Event Dinleme", level=2, color=GREEN)
code("# Time-lock başladı — UI'da sayaç göster")
code("filt = contract.events.TimeLockTriggered.create_filter(fromBlock='latest')")
code("for event in filt.get_new_entries():")
code("    tx_idx = event.args['txIndex']")
code("    unlock = event.args['unlockTime']")
code("    # UI'a gönder: 'İşlem {tx_idx} → {unlock} zamanında execute edilebilir'")
code("")
code("# Cüzdan donduruldu")
code("filt2 = contract.events.EmergencyPaused.create_filter(fromBlock='latest')")
code("for event in filt2.get_new_entries():")
code("    who = event.args['triggeredBy']")
code("    # UI'a gönder: '{who} cüzdanı dondurdu!'")

heading("9.5 Deploy Sonrası Yapman Gerekenler", level=2, color=GREEN)
para("Merve sana CONTRACT_ADDRESS verecek. .env dosyana ekle:", space_after=3)
code("CONTRACT_ADDRESS=0x...    # Merve'den alacaksın")
code("SEPOLIA_RPC_URL=...       # Infura/Alchemy URL")
code("CHAIN_ID=11155111         # Sepolia")
code("")
code("# Bağlantı doğrulama:")
code("w3 = Web3(Web3.HTTPProvider(os.getenv('SEPOLIA_RPC_URL')))")
code("assert w3.eth.chain_id == 11155111, 'Sepolia değil!'")
code("assert w3.is_connected(), 'Bağlantı yok!'")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  BÖLÜM 10 — İREM İÇİN: TEST VE UI
# ══════════════════════════════════════════════════════════════════════════════

heading("Bölüm 10 — İrem İçin: Test ve UI", level=1, color=ORANGE)
para(
    "Bu bölüm İrem'in test senaryoları ve UI geliştirmesi için "
    "ihtiyaç duyduğu her şeyi içeriyor.",
    size=10, color=MUTED, italic=True
)

heading("10.1 Tüm Revert Mesajları (Test İçin)", level=2, color=ORANGE)
para(
    "Kontrat belirli durumlarda işlemi durdurur ve bir hata mesajı verir. "
    "Testlerde bu mesajları birebir kullan:"
)
table(
    headers=["Senaryo", "Revert Mesajı"],
    rows=[
        ["Boş owner listesiyle deploy", '"MultiSigWallet: en az 1 sahip gerekli"'],
        ["M < N/2 ile deploy", '"MultiSigWallet: gecersiz M veya N degeri (M >= N/2 sarti)"'],
        ["address(0) owner", '"MultiSigWallet: sifir adres"'],
        ["Tekrar eden owner adresi", '"MultiSigWallet: sahip adresi tekrarlanamaz"'],
        ["Sahip olmayan submitTransaction çağırır", '"MultiSigWallet: sadece sahip cagirir"'],
        ["Hedef adres sıfır", '"MultiSigWallet: sifir hedef adres"'],
        ["Geçersiz txIndex", '"TransactionManager: islem bulunamadi"'],
        ["Yürütülmüş işlemi onaylamaya çalışmak", '"TransactionManager: islem zaten yurutuldu"'],
        ["Aynı sahip iki kez onaylar", '"TransactionManager: zaten onaylanmis"'],
        ["Onayı olmayan revoke", '"TransactionManager: onay bulunamadi"'],
        ["M onay olmadan execute", '"MultiSigWallet: yetersiz onay sayisi"'],
        ["onlyWallet fonksiyonunu direkt çağırmak", '"OwnerManager: sadece kontrat cagirabilir"'],
        ["Time-lock dolmadan execute", '"TransactionManager: time-lock suresi dolmadi"'],
        ["Paused iken herhangi bir işlem", '"MultiSigWallet: cuzcdan durduruldu"'],
        ["Zaten paused iken pause()", '"MultiSigWallet: zaten durduruldu"'],
        ["Paused değilken unpause()", '"MultiSigWallet: zaten aktif"'],
        ["Yanlış zincirde execute", '"MultiSigWallet: yanlis zincir"'],
        ["Execute başarısız (hedef hata verir)", '"MultiSigWallet: islem yurutme basarisiz"'],
    ],
    hdr_bg="854F0B",
    col_widths=[7, 11.5],
    row_size=8.5,
)

heading("10.2 Test Senaryoları", level=2, color=ORANGE)
table(
    headers=["Kategori", "Test", "Beklenen Sonuç"],
    rows=[
        ["2-of-3 akış",    "Submit → 2 onay → 24h bekle → execute",  "Başarılı — ETH transfer olur"],
        ["2-of-3 akış",    "Submit → 1 onay → execute",              "Revert: yetersiz onay"],
        ["Time-lock",      "M onay → anında execute",                 "Revert: time-lock dolmadı"],
        ["Time-lock",      "evm_increaseTime(86401) → execute",       "Başarılı"],
        ["Time-lock",      "TimeLockTriggered event emit edildi mi?",  "Event doğru parametrelerle geldi"],
        ["Pause",          "pause() → submitTransaction",              "Revert: cüzdan durduruldu"],
        ["Pause",          "pause() → unpause() → submitTransaction",  "Başarılı"],
        ["Pause",          "Sahip olmayan pause()",                   "Revert: sadece sahip çağırır"],
        ["Replay",         "Aynı txIndex iki kez execute",            "Revert: zaten yürütüldü"],
        ["Replay",         "Execute sonrası nonce +1 mi?",            "nonce == önceki + 1"],
        ["Reentrancy",     "Kötü amaçlı kontrat geri çağırır",        "Revert (ReentrancyGuard)"],
        ["Owner yönetimi", "addOwner direkt çağrılır",                "Revert: sadece kontrat çağırabilir"],
        ["Owner yönetimi", "removeOwner → M > N olursa",              "M otomatik düşürülür"],
    ],
    hdr_bg="854F0B",
    col_widths=[3.5, 7, 8],
)

heading("10.3 Test Ortamında Time-lock Hızlandırma", level=2, color=ORANGE)
para("24 saati gerçekten beklemek zorunda değilsin. Hardhat'ın time manipulation araçlarını kullan:", space_after=3)
code("// JavaScript test dosyasında:")
code("await network.provider.send('evm_increaseTime', [86401]); // 24 saat + 1 saniye")
code("await network.provider.send('evm_mine');                   // Yeni blok oluştur")
code("// Artık executeTransaction çalışır")

heading("10.4 UI İçin Kullanman Gereken Fonksiyonlar", level=2, color=ORANGE)
table(
    headers=["UI Bileşeni", "Contract Çağrısı", "Ne Döner?"],
    rows=[
        ["Pending tx listesi",       "getTransactionCount() + getTransaction(i)",    "Tüm işlemler"],
        ["İşlem onay durumu",        "getConfirmationStatus(txIndex, ownerAddr)",     "bool"],
        ["İşlem onaylanabilir mi?",  "isConfirmed(txIndex)",                         "bool"],
        ["ETH bakiyesi",             "getBalance()",                                  "uint256 (wei)"],
        ["Time-lock sayacı",         "readyTime(txIndex)",                            "unix timestamp"],
        ["Cüzdan dondurulmuş mu?",   "paused()",                                     "bool"],
        ["Sahip listesi",            "getOwners()",                                   "address[]"],
        ["Sahip mi?",                "isOwner(adres)",                               "bool"],
        ["Gereken onay sayısı",      "numConfirmationsRequired()",                    "uint256"],
    ],
    hdr_bg="854F0B",
    col_widths=[4.5, 6, 8],
)

heading("10.5 Canlı Güncelleme İçin Dinlemen Gereken Eventler", level=2, color=ORANGE)
table(
    headers=["Event", "Ne Zaman Gelir?", "UI'da Ne Yaparsın?"],
    rows=[
        ["SubmitTransaction",     "Yeni işlem teklif edilince",     "Pending listeyi güncelle"],
        ["ConfirmTransaction",    "Onay verilince",                 "Onay sayacını artır"],
        ["RevokeConfirmation",    "Onay geri çekilince",            "Onay sayacını azalt"],
        ["TimeLockTriggered",     "M onaya ulaşılınca",             "Geri sayım başlat"],
        ["ExecuteTransaction",    "İşlem yürütülünce",              "Pending'den kaldır"],
        ["EmergencyPaused",       "Cüzdan dondurulunca",            "Uyarı banner göster"],
        ["EmergencyUnpaused",     "Cüzdan açılınca",               "Uyarıyı kaldır"],
        ["Deposit",               "ETH yatırılınca",               "Bakiyeyi güncelle"],
    ],
    hdr_bg="854F0B",
    col_widths=[4.5, 5.5, 8.5],
)


# ── Kaydet ───────────────────────────────────────────────────────────────────

out = r"c:\Users\merve\Desktop\8. dönem\blockchain\blockchain-prohe\MultiSigWallet_Proje_Raporu.docx"
doc.save(out)
print(f"Rapor kaydedildi:\n{out}")
