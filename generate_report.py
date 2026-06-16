"""
MultiSig Wallet Proje Raporu - Word Belgesi Oluşturucu
COM 4532 Blockchain Technology and Public Ledgers
Ankara Universitesi | Merve Cakir, Emine Binay, Irem Keskin
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Sayfa kenar boslukları ──────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Yardimci fonksiyonlar ───────────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.5 + 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Baslik satiri
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3A5F')
        shd.set(qn('w:color'), 'FFFFFF')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Veri satirlari
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ════════════════════════════════════════════════════════════════
# KAPAK SAYFASI
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\nANKARA ÜNİVERSİTESİ\nMÜHENDİSLİK FAKÜLTESİ\nBİLGİSAYAR MÜHENDİSLİĞİ BÖLÜMÜ')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('COM 4532 — Blokzincir Teknolojisi ve Kamuya Açık Defterler')
run2.font.size = Pt(12)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('PROJE RAPORU')
run3.bold = True
run3.font.size = Pt(18)
run3.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('Çok İmzalı (Multi-Signature) Cüzdan Implementasyonu')
run4.bold = True
run4.font.size = Pt(15)

doc.add_paragraph()
doc.add_paragraph()

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run(
    'Merve ÇAKIR — 21290542\n'
    'Emine BİNAY — 22290370\n'
    'İrem KESKİN — 21290538'
)
run5.font.size = Pt(12)

doc.add_paragraph()

p6 = doc.add_paragraph()
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
run6 = p6.add_run('Dr. Murat OSMANOĞLU')
run6.font.size = Pt(11)

doc.add_paragraph()

p7 = doc.add_paragraph()
p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
run7 = p7.add_run(datetime.datetime.now().strftime('%d.%m.%Y'))
run7.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. PROJE ÖZETI
# ════════════════════════════════════════════════════════════════
heading('1. Proje Özeti', 1)

para(
    'Bu projede Ethereum blokzinciri üzerinde çalışan, Solidity ile yazılmış '
    'modüler bir çok imzalı (multi-signature) cüzdan sistemi geliştirilmiştir. '
    'Sistem; tek özel anahtara (private key) bağımlı olan tek imzalı cüzdanların '
    'güvenlik zafiyetini ortadan kaldırmayı hedefler. Bir işlemin yürütülebilmesi '
    'için N sahibin en az M tanesinin onayı zorunludur (M-of-N eşik şeması).',
    size=11
)

para(
    'Projemiz iki katmandan oluşmaktadır:',
    size=11
)
bullet('Zincir üstü (on-chain) katman: Solidity ile yazılmış 3 modüler akıllı kontrat — sahip yönetimi, işlem yaşam döngüsü ve güvenlik mekanizmaları.')
bullet('Zincir dışı (off-chain) katman: Web3.py ile blockchain etkileşimi sağlayan Python backend ve Streamlit ile yazılmış kullanıcı arayüzü.')

para(
    'Sistem Sepolia test ağına deploy edilmiş ve gerçek ETH transferleri gerçekleştirilmiştir.',
    size=11
)

divider()

# ════════════════════════════════════════════════════════════════
# 2. EKİP GÖREVLERİ
# ════════════════════════════════════════════════════════════════
heading('2. Ekip Görev Dağılımı', 1)

add_table(
    ['Kişi', 'Sorumluluk Alanı', 'Dosyalar'],
    [
        ['Merve Çakır', 'Akıllı kontrat tasarımı, güvenlik mekanizmaları, test altyapısı, proje koordinasyonu',
         'contracts/*.sol, test/test_multisig.py, test/gas_benchmark.py'],
        ['Emine Binay', 'Python backend, Web3.py entegrasyonu, blockchain bağlantısı',
         'irem/backend/wallet_service.py, irem/backend/config.py'],
        ['İrem Keskin', 'Streamlit kullanıcı arayüzü, oturum yönetimi, PDF raporu',
         'irem/app.py'],
    ],
    col_widths=[3.5, 7, 5.5]
)

doc.add_paragraph()
divider()

# ════════════════════════════════════════════════════════════════
# 3. SİSTEM MİMARİSİ
# ════════════════════════════════════════════════════════════════
heading('3. Sistem Mimarisi', 1)

para(
    '3.1 Akıllı Kontrat Katmanı',
    bold=True, size=12
)

para(
    'Proje üç birbirini miras alan (inheritance) Solidity kontratından oluşmaktadır. '
    'Bu modüler tasarım; saldırı yüzeyini küçültür, kod karmaşıklığını azaltır ve '
    'her bileşenin bağımsız olarak test edilmesini sağlar.',
    size=11
)

add_table(
    ['Kontrat', 'Görev', 'Kritik Özellikler'],
    [
        ['TransactionManager.sol',
         'İşlem yaşam döngüsü: submit → confirm → execute → revoke',
         'CEI pattern, reentrancy guard, time-lock, executed flag'],
        ['OwnerManager.sol',
         'Sahip yönetimi: ekleme, çıkarma, değiştirme, threshold güncelleme',
         'onlyWallet modifier — sadece multi-sig konsensüsüyle değişir'],
        ['MultiSigWallet.sol',
         'Ana kontrat — diğerlerini miras alır, emergency pause, DOMAIN_SEPARATOR',
         'EIP-712 replay koruması, pause/unpause, event yayınlama'],
    ],
    col_widths=[4, 6, 6]
)

doc.add_paragraph()
para('Kontrat Adresi (Sepolia Testnet):', bold=True, size=11)
code_block('0x5cD9338E23dec13b985DAE792E2D4E419B9443e4')
para('Ağ: Sepolia (chainId: 11155111)  |  Solidity: 0.8.20  |  OpenZeppelin: 5.x', size=10)

doc.add_paragraph()
para('3.2 Backend Katmanı (wallet_service.py)', bold=True, size=12)
para(
    'Web3.py 7.16.0 kullanılarak yazılan Python backend, akıllı kontratla tüm '
    'etkileşimi yönetir. Her fonksiyon private key parametresi alarak farklı '
    'sahiplerin kendi kimliğiyle işlem imzalamasına izin verir.',
    size=11
)

add_table(
    ['Fonksiyon', 'Açıklama'],
    [
        ['get_wallet_stats()', 'Bakiye, threshold, sahip listesi, kontrat adresi'],
        ['get_pending_transactions()', 'Yürütülmemiş işlemlerin listesi'],
        ['submit_transaction(to, amount, data)', 'Yeni işlem teklifi oluşturur'],
        ['approve_transaction(tx_id)', 'İşlemi onaylar (confirmTransaction)'],
        ['execute_transaction(tx_id)', 'Onaylanmış işlemi yürütür'],
        ['revoke_confirmation(tx_id)', 'Verilen onayı geri çeker'],
        ['pause_wallet() / unpause_wallet()', 'Acil durdurma mekanizması'],
        ['encode_management_call(fn, *args)', 'addOwner/removeOwner işlemlerini ABI-encode eder'],
        ['get_paused_status()', 'Cüzdanın dondurulup dondurulmadığını döner'],
    ],
    col_widths=[6, 10]
)

doc.add_paragraph()
para('3.3 Frontend Katmanı (Streamlit Dashboard)', bold=True, size=12)
para(
    'Streamlit 1.50.0 ile yazılmış kullanıcı arayüzü; private key ve kontrat adresi '
    'ile giriş yapılmasını sağlar. Oturum bilgisi .session.json dosyasına kaydedilir, '
    'sayfa yenilendiğinde kullanıcı tekrar giriş yapmak zorunda kalmaz.',
    size=11
)

bullet('Dashboard metrikleri: bakiye, bekleyen işlem sayısı, toplam sahip')
bullet('Pending Approvals sekmesi: onay durumu pasta grafiği, onay/yürüt butonları')
bullet('Action Center: yeni transfer önerisi, owner yönetimi (addOwner/removeOwner/replaceOwner/changeThreshold)')
bullet('Sidebar: aktif hesap gösterimi, Emergency Pause/Unpause butonu, PDF raporu indirme, çıkış')
bullet('Cüzdan dondurulduğunda tüm action butonları devre dışı kalır')

divider()

# ════════════════════════════════════════════════════════════════
# 4. GÜVENLİK MEKANİZMALARI
# ════════════════════════════════════════════════════════════════
heading('4. Güvenlik Mekanizmaları ve Saldırı Vektörleri', 1)

para(
    'Proposal\'da tanımlanan 5 saldırı vektörünün tamamına karşı önlem alınmıştır.',
    size=11
)

add_table(
    ['Saldırı Vektörü', 'Tehdit', 'Uygulanan Önlem', 'Test'],
    [
        ['Signature Replay',
         'Geçmiş bir imzanın farklı TX için tekrar kullanılması',
         'Her TX\'e benzersiz txIndex atanır; kontrat adresi ve chainId (EIP-155) hash\'e dahil edilir',
         'test_replay_attack_tx0_approval_does_not_apply_to_tx1'],
        ['Reentrancy (Tekrar Giriş)',
         'Dış kontratın para çekerken tekrar withdraw çağırması',
         'CEI (Checks-Effects-Interactions) pattern + OpenZeppelin ReentrancyGuard mutex kilidi',
         'test_reentrancy_guard_executed_flag'],
        ['Front-Running / MEV',
         'Mempool\'daki TX\'i gören bots\'un öne geçmesi',
         'Time-lock mekanizması — onaydan sonra bekleme süresi zorunlu; changeTimeLockDuration ile ayarlanabilir',
         'Kontrat seviyesinde doğrulandı'],
        ['Social Engineering',
         'Sahte arayüz ile yanlış TX onaylatma',
         'Private key ile doğrudan kontrat imzalama, TX detayları arayüzde gösterilir',
         'Manuel doğrulama'],
        ['Quorum Manipulation',
         'Sahiplerin anlaşarak threshold\'u düşürmesi',
         'Threshold değişikliği sadece onlyWallet modifier\'ı ile çalışır — mevcut M imzası gerektirir',
         'test_unauthorized_execute_fails'],
        ['Emergency Pause',
         'Açık keşfedildiğinde işlemlerin durdurulması gerekebilir',
         'pause()/unpause() fonksiyonları; whenNotPaused modifier tüm TX fonksiyonlarında',
         'TestEmergencyPause sınıfı (4 test)'],
    ],
    col_widths=[3.5, 4, 5, 3.5]
)

doc.add_paragraph()
para('4.1 OpenZeppelin ReentrancyGuard — Nasıl Çalışır?', bold=True, size=11)
para(
    'executeTransaction çağrıldığında önce executed bayrağı true yapılır, '
    'ardından ETH transferi gerçekleşir. Kötü niyetli bir kontrat receive() '
    'fonksiyonundan tekrar executeTransaction çağırsa dahi executed == true '
    'kontrolü bunu reddeder. ReentrancyGuard ek olarak nonReentrant modifier '
    'ile fonksiyon girişinde kilit koyar.',
    size=11
)

para('4.2 EIP-155 Replay Koruması — Nasıl Çalışır?', bold=True, size=11)
para(
    'Her işlem kendi txIndex numarasını taşır. TX0\'a verilen onay yalnızca '
    'TX0 için geçerlidir; TX1\'e geçmez. Ayrıca DOMAIN_SEPARATOR içinde '
    'kontrat adresi ve chainId (11155111) yer alır; başka bir zincirde veya '
    'başka bir kontrat adresinde imza geçersiz sayılır.',
    size=11
)

divider()

# ════════════════════════════════════════════════════════════════
# 5. STATİK ANALİZ — SLİTHER
# ════════════════════════════════════════════════════════════════
heading('5. Statik Analiz — Slither', 1)

para(
    'Slither v0.11.4 ile 5 kontrat, 100 dedektör kullanılarak analiz edilmiştir. '
    'Hiçbir kritik (Critical) veya yüksek (High) seviyeli güvenlik açığı bulunamamıştır.',
    size=11
)

code_block('Komut: slither contracts/MultiSigWallet.sol --solc-remaps "@openzeppelin/contracts/=node_modules/@openzeppelin/contracts/"')
code_block('Sonuc: 5 contracts with 100 detectors — 28 result(s) found')
code_block('Critical: 0  |  High: 0  |  Medium: 0')

doc.add_paragraph()

add_table(
    ['Bulgu Türü', 'Nerede', 'Açıklama', 'Önem'],
    [
        ['assembly', 'OpenZeppelin StorageSlot.sol', 'OZ kütüphanesi iç assembly — bizim kodumuzdaki değil', 'Informational'],
        ['pragma', 'OZ vs kontratlarımız', 'OZ ^0.8.20 kullanıyor, bizim 0.8.20 — zararsız', 'Low'],
        ['solc-version', 'Tüm kontratlar', '0.8.20\'de bilinen küçük bug\'lar (VerbatimInvalidDeduplication vb.) — pratikte etkisiz', 'Low'],
        ['low-level-calls', 'executeTransaction()', '.call{value}() kullanımı — multi-sig için beklenen, CEI+Guard ile korunuyor', 'Informational'],
        ['naming-convention', 'Tüm parametreler', '_to, _value gibi underscore prefix — stil sorunu, güvenlik değil', 'Informational'],
    ],
    col_widths=[3.5, 4, 7.5, 2.5]
)

divider()

# ════════════════════════════════════════════════════════════════
# 6. TEST SONUÇLARI
# ════════════════════════════════════════════════════════════════
heading('6. Test Sonuçları — 33/33 PASSED', 1)

para(
    'Hardhat local node üzerinde çalıştırılan pytest test paketi, her test için '
    'bağımsız bir kontrat deploy edip 1 ETH yüklemektedir. 33 testin tamamı geçmektedir.',
    size=11
)

code_block('Calistirma: npx hardhat node  (Terminal 1)')
code_block('             python -m pytest test/test_multisig.py -v  (Terminal 2)')
code_block('Sonuc: 33 passed in 3.55s')

doc.add_paragraph()

add_table(
    ['Test Sınıfı', 'Test Sayısı', 'Ne Test Eder'],
    [
        ['TestDeploy', '7', 'Doğru owner listesi, threshold, bakiye; boş/hatalı parametrelerle deploy reddi, duplicate owner reddi'],
        ['TestSubmitTransaction', '4', 'Owner TX önerebilir, non-owner öneremez, TX sayısı artar, SubmitTransaction eventi yayınlanır'],
        ['TestConfirmTransaction', '5', 'Owner onay verebilir, non-owner veremez, çift onay reddedilir, geçersiz TX reddedilir, onay sayısı artar'],
        ['TestExecuteTransaction', '5', 'Yeterli onayda execute başarılı, alıcı ETH alır, çift execute reddedilir, yetersiz onayda reddedilir, executed=true'],
        ['TestRevokeConfirmation', '4', 'Owner revoke edebilir, onay vermemiş owner edemez, revoke sonrası execute reddedilir, sayı azalır'],
        ['TestSecurityAttacks', '4', 'Replay koruması, yetkisiz execute reddi, reentrancy guard (executed flag), deposit bakiye artışı'],
        ['TestEmergencyPause', '4', 'Owner dondurabilir, dondurulunca submit reddedilir, owner açabilir ve işlem tekrar çalışır, non-owner donduramaz'],
    ],
    col_widths=[4.5, 2, 9.5]
)

divider()

# ════════════════════════════════════════════════════════════════
# 7. GAS VERİMLİLİĞİ ANALİZİ
# ════════════════════════════════════════════════════════════════
heading('7. Gas Verimliliği Analizi', 1)

para(
    'Hardhat local node üzerinde her işlem adımı için gas tüketimi ölçülmüş ve '
    'Gnosis Safe v1.3.0 referans değerleriyle karşılaştırılmıştır.',
    size=11
)

code_block('Calistirma: python test/gas_benchmark.py')

doc.add_paragraph()

add_table(
    ['İşlem', 'MultiSigWallet (Bizim)', 'Gnosis Safe v1.3.0', 'Fark (%)'],
    [
        ['Deploy', '2,232,447', '1,349,852', '+65.4%'],
        ['submitTransaction', '106,188', '72,400', '+46.7%'],
        ['confirmTransaction (1. onay)', '82,231', '48,200', '+70.6%'],
        ['confirmTransaction (2. onay)', '113,349', '48,200', '+135.2%'],
        ['executeTransaction', '101,704', '65,900', '+54.3%'],
        ['TX Lifecycle Toplam', '403,472', '234,700', '+71.9%'],
    ],
    col_widths=[5, 4, 4, 3]
)

doc.add_paragraph()
para('Fark Analizi:', bold=True, size=11)
bullet('Gnosis Safe proxy mimarisi (EIP-1167 minimal proxy) kullanır — deploy gas\'ı düşüktür.')
bullet('Bizim kontratımız tek parça (monolithic) deploy edilir; bu akademik prototip için beklenen bir sonuçtur.')
bullet('2. confirmTransaction daha fazla gas harcar çünkü bu onay threshold\'u tamamlar ve isConfirmed döngüsü çalışır.')
bullet('Üretim sisteminde EIP-712 off-chain imza toplama ile confirmTransaction gas\'ı sıfıra indirilebilir.')

divider()

# ════════════════════════════════════════════════════════════════
# 8. İŞLEM AKIŞI — NASIL ÇALIŞIR
# ════════════════════════════════════════════════════════════════
heading('8. İşlem Akışı — Adım Adım', 1)

para('8.1 Normal Para Transferi (2-of-3)', bold=True, size=12)

steps = [
    ('Adım 1 — Submit', 'Bir sahip Streamlit\'te "Propose New Transfer" bölümünden alıcı adres ve miktar girer. submit_transaction() çağrılır, kontrata yazılır, txIndex alınır. Diğer sahipler bunu "Pending Approvals" tablosunda görür.'),
    ('Adım 2 — Confirm', 'İkinci sahip (farklı private key ile giriş yapmış) Pending listesinden TX\'i seçer, "Confirm Approval" butonuna basar. confirmTransaction() çağrılır, onay sayısı 1 artar.'),
    ('Adım 3 — Threshold', 'Threshold 2 olduğundan 2. onay yeterlidir. Pending tabloda "Ready" olarak işaretlenir, Execute butonu aktif hale gelir.'),
    ('Adım 4 — Execute', 'Herhangi bir sahip "Execute Transaction" butonuna basar. executeTransaction() çağrılır. Kontrat ETH\'i alıcıya gönderir. executed=true set edilir, tekrar yürütülemez.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(title + ': ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

doc.add_paragraph()
para('8.2 Owner Yönetimi (addOwner / removeOwner)', bold=True, size=12)
para(
    'Owner değişiklikleri doğrudan yapılamaz — mutlaka mevcut sahiplerin multi-sig '
    'konsensüsünden geçmesi gerekir. Bu mekanizma quorum manipulation saldırısını önler.',
    size=11
)

steps2 = [
    ('Adım 1', 'Bir sahip "Owner Yönetimi" panelinden işlem türünü (Add/Remove/Replace) ve adresi seçer.'),
    ('Adım 2', 'encode_management_call() fonksiyonu, kontrat fonksiyonunu ABI-encode eder (addOwner(0x...) calldata\'sı).'),
    ('Adım 3', 'Bu calldata, TO adresi kontratın kendisi olarak submitTransaction\'a verilir.'),
    ('Adım 4', 'Diğer sahipler normal akıştaki gibi onaylar. Threshold dolunca execute edildiğinde kontrat kendi addOwner fonksiyonunu çalıştırır.'),
]
for title, desc in steps2:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(title + ': ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

doc.add_paragraph()
para('8.3 Emergency Pause', bold=True, size=12)
para(
    'Herhangi bir sahip sidebar\'dan "Emergency Pause" butonuna basarak cüzdanı '
    'anlık olarak dondurabilir. Dondurulduğunda submitTransaction, confirmTransaction, '
    'executeTransaction ve revokeConfirmation fonksiyonlarının tamamı whenNotPaused '
    'modifier\'ı ile reddedilir. Arayüzde Action Center komple kapanır ve kırmızı '
    'uyarı banner\'ı gösterilir. Herhangi bir sahip "Yeniden Aç" ile kaldırabilir.',
    size=11
)

divider()

# ════════════════════════════════════════════════════════════════
# 9. TEKNİK DETAYLAR
# ════════════════════════════════════════════════════════════════
heading('9. Teknik Detaylar', 1)

para('9.1 Kullanılan Teknoloji ve Sürümler', bold=True, size=12)

add_table(
    ['Teknoloji', 'Sürüm', 'Kullanım Amacı'],
    [
        ['Solidity', '0.8.20', 'Akıllı kontrat programlama dili'],
        ['OpenZeppelin Contracts', '5.x', 'ReentrancyGuard, güvenlik kütüphaneleri'],
        ['Hardhat', '3.x (ESM)', 'Lokal blockchain, derleme, deploy'],
        ['Web3.py', '7.16.0', 'Python-blockchain etkileşimi'],
        ['Streamlit', '1.50.0', 'Web kullanıcı arayüzü'],
        ['Python', '3.9', 'Backend dili'],
        ['Slither', '0.11.4', 'Statik güvenlik analizi'],
        ['pytest', '8.4.2', 'Otomatik test çerçevesi'],
        ['fpdf2', '2.8.4', 'PDF raporu üretimi'],
        ['Sepolia Testnet', 'chainId: 11155111', 'Test ağı deploy'],
        ['Infura', '-', 'RPC sağlayıcısı'],
    ],
    col_widths=[4, 3, 9]
)

doc.add_paragraph()
para('9.2 Kritik Kod Kararları', bold=True, size=12)

bullets = [
    ('web3.py v7\'de signed.raw_transaction (rawTransaction değil)',
     'Web3.py 7.x API değişikliği — eski sürümden farklı.'),
    ('bytes(pdf.output()) — fpdf2\'de',
     'fpdf2 kütüphanesi bytearray döner, str değil. encode() çağrısı hata verir.'),
    ('st.cache_data(ttl=8) — blockchain çağrıları için',
     'Her render\'da RPC çağrısı yapmak uygulamayı yavaşlatır. 8 saniyelik cache yeterince güncel tutar.'),
    ('gasPrice: w3.eth.gas_price — test tx\'lerinde',
     'EIP-1559 desteklemeyen legacy tx formatı Hardhat local node\'unda gasPrice gerektirir.'),
    ('Web3.to_checksum_address() — tüm adreslerde',
     'web3.py v7 büyük/küçük harf duyarlıdır; Hardhat\'ın default adresleri checksum olmadan InvalidAddress hatası verir.'),
    ('onlyWallet modifier',
     'addOwner/removeOwner gibi admin fonksiyonları sadece msg.sender == address(this) ise çalışır. Yani ancak kontratın kendisi çağırabilir — çok imzalı konsensüs zorunludur.'),
]

for title, desc in bullets:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run('• ' + title + ': ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10.5)

divider()

# ════════════════════════════════════════════════════════════════
# 10. PROPOSAL KARŞILAŞTIRMASı
# ════════════════════════════════════════════════════════════════
heading('10. Proposal Vaadleri vs. Gerçekleştirilenler', 1)

add_table(
    ['Vaad (Proposal)', 'Durum', 'Açıklama'],
    [
        ['3 modüler Solidity kontrat', 'YAPILDI', 'MultiSigWallet, OwnerManager, TransactionManager'],
        ['OpenZeppelin ReentrancyGuard', 'YAPILDI', 'nonReentrant + CEI pattern uygulandı'],
        ['EIP-155 nonce-based replay koruması', 'YAPILDI', 'txIndex + chainId + contract address'],
        ['Python + Web3.py backend', 'YAPILDI', 'wallet_service.py — 12 fonksiyon'],
        ['Streamlit dashboard', 'YAPILDI', 'Login, session persistence, tüm işlemler UI\'dan'],
        ['Sepolia testnet deploy', 'YAPILDI', '0x5cD9338E23dec13b985DAE792E2D4E419B9443e4'],
        ['emergencyPause mekanizması', 'YAPILDI', 'pause()/unpause() + whenNotPaused + UI butonu'],
        ['Time-lock mekanizması', 'YAPILDI', 'Kontrakta var; demo için 0s, changeTimeLockDuration ile ayarlanabilir'],
        ['addOwner/removeOwner UI', 'YAPILDI', 'Owner Yönetimi paneli — multi-sig konsensüsüyle'],
        ['Slither statik analiz', 'YAPILDI', '0 kritik, 0 yüksek seviye açık'],
        ['Gas analizi — Gnosis Safe karşılaştırması', 'YAPILDI', 'benchmark script, tablo halinde raporlandı'],
        ['Otomatik test paketi', 'YAPILDI', '33 test, 7 sınıf, tamamı geçiyor'],
        ['Saldırı simülasyon testleri', 'YAPILDI', 'Replay, Reentrancy, Quorum, Pause testleri'],
        ['Mythril sembolik analiz', 'ATILDI', 'Docker bağımlılığı; Slither kapsamlı sonuçlar verdi'],
        ['%90+ code coverage ölçümü', 'ATILDI', 'Hardhat 3.x ile solidity-coverage uyumsuz; 33 test tüm public fonksiyonları kapsıyor'],
        ['EIP-712 off-chain imza toplama', 'ATILDI', 'Karmaşıklık nedeniyle on-chain yaklaşım tercih edildi; gas farkı analiz raporunda gösterildi'],
    ],
    col_widths=[6, 2.5, 7.5]
)

divider()

# ════════════════════════════════════════════════════════════════
# 11. ÇALIŞTIRILMASI
# ════════════════════════════════════════════════════════════════
heading('11. Projenin Çalıştırılması', 1)

para('Gereksinimler:', bold=True, size=11)
bullet('Node.js v18+, Python 3.9+')
bullet('npm install  (proje kökünde)')
bullet('pip install -r irem/requirements.txt')

doc.add_paragraph()
para('Adım 1 — Lokal test (Hardhat node):', bold=True, size=11)
code_block('Terminal 1:  npx hardhat node')
code_block('Terminal 2:  python -m pytest test/test_multisig.py -v')
code_block('             python test/gas_benchmark.py')

doc.add_paragraph()
para('Adım 2 — Streamlit arayüzü (Sepolia):', bold=True, size=11)
code_block('cd irem')
code_block('streamlit run app.py')
code_block('Tarayici: http://localhost:8501')
code_block('Private Key ve Kontrat Adresi ile giris yapilir.')

doc.add_paragraph()
para('Adım 3 — Statik analiz:', bold=True, size=11)
code_block('slither contracts/MultiSigWallet.sol --solc-remaps "@openzeppelin/contracts/=node_modules/@openzeppelin/contracts/"')

divider()

# ── FOOTER ──────────────────────────────────────────────────────
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_footer = p_footer.add_run(
    f'COM 4532 Blockchain Technology and Public Ledgers  |  Ankara Üniversitesi  |  {datetime.datetime.now().strftime("%d.%m.%Y")}'
)
run_footer.font.size = Pt(9)
run_footer.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── KAYDET ──────────────────────────────────────────────────────
output_path = r'c:\Users\merve\Desktop\8. dönem\blockchain\blockchain-prohe\COM4532_MultiSig_Proje_Raporu.docx'
doc.save(output_path)
print(f'Rapor olusturuldu: {output_path}')
